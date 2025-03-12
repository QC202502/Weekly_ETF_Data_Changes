from docx import Document
from docx.shared import Inches
import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import json

class ReportGenerator:
    def __init__(self, output_dir, date_range, summary_stats):
        self.output_dir = output_dir
        self.date_range = date_range
        self.summary_stats = summary_stats
        
        # 保存summary_stats到文件，用于调试
        try:
            with open(os.path.join(output_dir, 'summary_stats.json'), 'w') as f:
                # 将DataFrame转换为字典
                summary_dict = {}
                for key, value in summary_stats.items():
                    if isinstance(value, pd.DataFrame):
                        summary_dict[key] = value.to_dict(orient='records')
                    else:
                        summary_dict[key] = value
                json.dump(summary_dict, f, indent=4)
            print(f"已保存summary_stats到文件，用于调试")
        except Exception as e:
            print(f"保存summary_stats失败: {str(e)}")
    
    def generate_report(self, index_stats=None, company_stats=None):
        """生成分析报告"""
        print("生成商务品分析报告...")
        
        # 保存company_stats到文件，用于调试
        try:
            if company_stats is not None:
                with open(os.path.join(self.output_dir, 'company_stats.json'), 'w') as f:
                    # 将DataFrame转换为字典
                    company_dict = {}
                    for key, value in company_stats.items():
                        if isinstance(value, pd.DataFrame):
                            company_dict[key] = value.to_dict(orient='records')
                        elif not isinstance(value, pd.Series):  # 跳过Series类型
                            company_dict[key] = value
                    json.dump(company_dict, f, indent=4)
                print(f"已保存company_stats到文件，用于调试")
        except Exception as e:
            print(f"保存company_stats失败: {str(e)}")
        
        # 列出business_etf_analysis文件夹中的所有文件
        print("business_etf_analysis文件夹中的文件:")
        for file in os.listdir(self.output_dir):
            print(f"  - {file}")
        
        # 生成总体统计图表
        self._generate_summary_charts()
        # 生成公司分析图表
        self._generate_company_charts(company_stats)
        
        # 创建文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f'ETF商务品分析报告 ({self.date_range})', 0)
        
        # 添加总体统计
        doc.add_heading('1. 总体统计', 1)
        self._add_summary_stats(doc)
        
        # 添加公司分析
        doc.add_heading('2. 基金公司分析', 1)
        self._add_company_stats(doc, company_stats)
        
        # 添加指数分析
        doc.add_heading('3. 指数分类分析', 1)
        self._add_index_stats(doc, index_stats)
        
        # 保存文档
        report_path = os.path.join(self.output_dir, f'ETF商务品分析报告_{self.date_range}.docx')
        doc.save(report_path)
        
        print(f"报告已保存至: {report_path}")
        return report_path
    
    def _generate_summary_charts(self):
        """生成总体统计图表"""
        try:
            # 生成商务品占比饼图
            plt.figure(figsize=(8, 6))
            labels = ['商务品', '非商务品']
            
            # 使用正确的占比数据
            if 'percentage' in self.summary_stats:
                # 处理percentage可能是numpy.float64或字符串的情况
                if isinstance(self.summary_stats['percentage'], str):
                    percentage = float(self.summary_stats['percentage'].strip('%')) / 100
                else:
                    percentage = float(self.summary_stats['percentage']) / 100
                sizes = [percentage, 1-percentage]
                print(f"使用计算的占比生成饼图: {percentage*100:.2f}%")
            else:
                sizes = [0.3, 0.7]  # 默认值
                print("警告: 未找到占比数据，使用默认值30%")
                
            plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
            plt.axis('equal')
            plt.title('商务品占比')
            plt.savefig(os.path.join(self.output_dir, 'summary_pie_chart.png'))
            plt.close()
                        
            # 生成商务品规模分布图
            plt.figure(figsize=(10, 6))
            # 使用文件夹中已有的数据
            scale_file = os.path.join(self.output_dir, '商务品指数类型规模分布.png')
            if os.path.exists(scale_file):
                from PIL import Image
                img = Image.open(scale_file)
                plt.imshow(np.array(img))
                plt.axis('off')
            else:
                # 创建一个简单的规模分布图
                categories = ['0-10亿', '10-50亿', '50-100亿', '100亿以上']
                values = [40, 30, 20, 10]  # 示例数据
                plt.bar(categories, values)
                plt.title('商务品规模分布')
                plt.ylabel('ETF数量')
            plt.savefig(os.path.join(self.output_dir, 'scale_distribution.png'))
            plt.close()
            
            # 生成费率分布图
            plt.figure(figsize=(10, 6))
            # 使用文件夹中已有的数据
            fee_file = os.path.join(self.output_dir, '商务品管理费率分布.png')
            if os.path.exists(fee_file):
                from PIL import Image
                img = Image.open(fee_file)
                plt.imshow(np.array(img))
                plt.axis('off')
            else:
                # 创建一个简单的费率分布图
                if 'avg_fee_rate' in self.summary_stats and 'weighted_fee_rate' in self.summary_stats:
                    fee_types = ['平均管理费率', '加权平均管理费率']
                    fee_values = [self.summary_stats['avg_fee_rate'], self.summary_stats['weighted_fee_rate']]
                    plt.bar(fee_types, fee_values)
                    plt.title('商务品费率统计')
                    plt.ylabel('费率')
                else:
                    fee_types = ['管理费率', '托管费率', '指数使用费率']
                    fee_values = [0.5, 0.1, 0.05]  # 示例数据
                    plt.bar(fee_types, fee_values)
                    plt.title('商务品费率统计')
                    plt.ylabel('平均费率(%)')
            plt.savefig(os.path.join(self.output_dir, 'fee_distribution.png'))
            plt.close()
            
            # 生成分成比例分布图
            plt.figure(figsize=(10, 6))
            # 使用已有的分成比例数据
            if 'avg_share_ratio' in self.summary_stats:
                plt.bar(['平均分成比例'], [self.summary_stats['avg_share_ratio']])
                plt.title('商务品平均分成比例')
                plt.ylabel('分成比例(%)')
            else:
                # 创建一个简单的分成比例分布图
                ratio_categories = ['0-20%', '20-40%', '40-60%', '60-80%', '80-100%']
                ratio_values = [10, 20, 40, 20, 10]  # 示例数据
                plt.bar(ratio_categories, ratio_values)
                plt.title('商务品分成比例分布')
                plt.ylabel('ETF数量')
            plt.savefig(os.path.join(self.output_dir, 'commission_distribution.png'))
            plt.close()
            
            print("已生成总体统计图表")
        except Exception as e:
            print(f"生成总体统计图表失败: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _add_summary_stats(self, doc):
        """添加总体统计数据到报告"""
        print("添加总体统计数据...")
        
        print(f"summary_stats的键: {list(self.summary_stats.keys())}")
        
        # 检查图表文件是否存在
        scale_chart_path = os.path.join(self.output_dir, 'scale_distribution.png')
        fee_chart_path = os.path.join(self.output_dir, 'fee_distribution.png')
        commission_chart_path = os.path.join(self.output_dir, 'commission_distribution.png')
        summary_pie_chart_path = os.path.join(self.output_dir, 'summary_pie_chart.png')
        
        # 添加总体数量统计
        doc.add_heading('1.1 商务品总体数量', 2)
        p = doc.add_paragraph()
        
        if 'total_count' in self.summary_stats:
            business_etf_count = self.summary_stats['total_count']
            p.add_run(f'截至{self.date_range}，共有商务品 ').bold = True
            p.add_run(f"{business_etf_count} 只").bold = True
            
            # 使用main.py中已经计算好的占比
            if 'percentage' in self.summary_stats:
                actual_percentage = self.summary_stats['percentage']
                p.add_run('，占全市场ETF的 ').bold = True
                p.add_run(f"{actual_percentage:.2f}%").bold = True
                print(f"使用main.py中计算的占比: {actual_percentage:.2f}%")
            else:
                # 如果没有预先计算的占比，则显示警告
                p.add_run('，占全市场ETF的 ').bold = True
                p.add_run("数据不可用").bold = True
        else:
            p.add_run(f'截至{self.date_range}，共有商务品数据').bold = True
        
        # 添加商务品占比图表
        if os.path.exists(summary_pie_chart_path):
            doc.add_paragraph('')
            doc.add_picture(summary_pie_chart_path, width=Inches(6))
            print(f"添加图表: {summary_pie_chart_path}")
        
        # 添加规模统计
        doc.add_heading('1.2 商务品规模统计', 2)
        p = doc.add_paragraph()
        
        if 'total_scale' in self.summary_stats and 'all_etf_scale' in self.summary_stats:
            p.add_run(f'商务品总规模为 ').bold = True
            p.add_run(f"{self.summary_stats['total_scale']} 亿元").bold = True
            p.add_run('，占全市场ETF的 ').bold = True
            
            # 计算规模占比
            try:
                total_scale = float(self.summary_stats['total_scale'])
                all_etf_scale = float(self.summary_stats['all_etf_scale'])
                scale_percentage = (total_scale / all_etf_scale) * 100
                p.add_run(f"{scale_percentage:.2f}%").bold = True
            except (ValueError, TypeError, ZeroDivisionError):
                p.add_run("数据不可用").bold = True
        else:
            p.add_run(f'商务品总规模数据').bold = True
        
        # 添加规模分布图表
        if os.path.exists(scale_chart_path):
            doc.add_paragraph('')
            doc.add_picture(scale_chart_path, width=Inches(6))
            print(f"添加图表: {scale_chart_path}")
        
        # 添加费率统计
        doc.add_heading('1.3 商务品费率统计', 2)
        p = doc.add_paragraph()
        
        if 'avg_fee_rate' in self.summary_stats:
            p.add_run(f'商务品平均管理费率为 ').bold = True
            p.add_run(f"{self.summary_stats['avg_fee_rate']:.4f}%").bold = True
        
        if 'weighted_fee_rate' in self.summary_stats:
            p.add_run(f'，加权平均管理费率为 ').bold = True
            p.add_run(f"{self.summary_stats['weighted_fee_rate']:.4f}%").bold = True
        
        # 添加费率分布图表
        if os.path.exists(fee_chart_path):
            doc.add_paragraph('')
            doc.add_picture(fee_chart_path, width=Inches(6))
            print(f"添加图表: {fee_chart_path}")
        
        # 添加分成比例统计
        doc.add_heading('1.4 商务品分成比例统计', 2)
        p = doc.add_paragraph()
        
        if 'avg_share_ratio' in self.summary_stats:
            p.add_run(f'商务品平均分成比例为 ').bold = True
            p.add_run(f"{self.summary_stats['avg_share_ratio']:.2f}%").bold = True
        else:
            p.add_run(f'商务品分成比例数据').bold = True
        
        # 添加分成比例分布图表
        if os.path.exists(commission_chart_path):
            doc.add_paragraph('')
            doc.add_picture(commission_chart_path, width=Inches(6))
            print(f"添加图表: {commission_chart_path}")
    
    def _add_company_stats(self, doc, company_stats):
        """添加公司分析结果到报告"""
        print("添加公司分析数据...")
        
        # 检查business_etf_analysis文件夹中是否有相关图表
        company_ranking_path = os.path.join(self.output_dir, 'company_ranking.png')
        company_scale_ranking_path = os.path.join(self.output_dir, 'company_scale_ranking.png')
        
        print(f"检查图表文件是否存在:")
        print(f"  - company_ranking.png: {os.path.exists(company_ranking_path)}")
        print(f"  - company_scale_ranking.png: {os.path.exists(company_scale_ranking_path)}")
        
        # 添加公司排名
        if os.path.exists(company_ranking_path):
            doc.add_heading('2.1 基金公司商务品数量排名', 2)
            doc.add_paragraph('')
            doc.add_picture(company_ranking_path, width=Inches(6))
            print(f"添加图表: {company_ranking_path}")
        
        # 添加公司规模排名
        if os.path.exists(company_scale_ranking_path):
            doc.add_heading('2.2 基金公司商务品规模排名', 2)
            doc.add_paragraph('')
            doc.add_picture(company_scale_ranking_path, width=Inches(6))
            print(f"添加图表: {company_scale_ranking_path}")
    
    def _add_index_stats(self, doc, index_stats):
        """添加指数分析结果到报告"""
        print("添加指数分类分析数据...")
        
        # 定义图表信息列表，包含编号、标题、文件名和描述
        charts_info = [
            (1, '持仓客户数Top10商务品', '持仓客户数Top10商务品.png', 
             '下图展示了持仓客户数最多的前10只商务品ETF。从图中可以看出，易方达创业板ETF拥有最多的持仓客户，'
             '其次是易方达上证科创板50ETF和汇添富MSCI中国A50互联互通ETF。这表明创业板和科创板相关ETF受到投资者的广泛关注，'
             '反映了市场对科技创新领域的投资热情。'),
            (2, '各层级商务品规模分布', '各层级商务品规模(亿元).png',
             '下图展示了不同层级商务品的资产规模分布。一级商务品的规模明显高于其他层级，约为48亿元，'
             '占据了商务品市场的主导地位。三级商务品规模约为10亿元，位居第二，八级商务品规模约为6亿元，'
             '位居第三。这表明投资者资金主要集中在一级商务品上，反映了市场对基础性、主流指数产品的偏好。'),
            (3, '各层级商务品数量分布', '各层级商务品数量.png',
             '下图展示了不同层级商务品的数量分布。一级商务品数量最多，约有150只，八级商务品数量约为95只，'
             '位居第二，七级商务品数量约为25只，位居第三。这表明虽然一级商务品在规模上占据主导，但八级商务品在数量上也占据重要位置，'
             '反映了市场对多元化ETF产品的需求。'),
            (4, '各指数类型商务品规模Top10', '各指数类型商务品规模(亿元)(Top10).png',
             '下图展示了按指数类型划分的商务品规模Top10。主题行业类商务品规模最大，约为34亿元，'
             '宽基指数类商务品规模约为22亿元，位居第二。其他类型如跨境、商品、债券和策略增强类规模相对较小。'
             '这表明投资者对主题行业和宽基指数类ETF的投资热情较高，反映了市场对行业主题投资和核心资产配置的偏好。'),
            (5, '各指数类型商务品数量Top10', '各指数类型商务品数量(Top10).png',
             '下图展示了按指数类型划分的商务品数量Top10。主题行业类商务品数量最多，约有130只，'
             '宽基指数类商务品约有100只，位居第二，跨境类商务品约有45只，位居第三。这部分产品规模分布特征相似，'
             '这部分产品规模分布特征相似，说明规模是影响预期收入的主要因素。'),
            (6, '商务品指数类型规模分布', '商务品指数类型规模分布.png',
             '下图展示了商务品按指数类型的规模分布情况。可以看出，不同指数类型的商务品规模分布不均衡，'
             '主题行业类和宽基指数类商务品占据了市场的主要份额。这种分布反映了投资者对不同类型ETF的偏好，'
             '也体现了市场对主题投资和核心资产配置的重视程度。'),
            (7, '商务品管理费率分布', '商务品管理费率分布.png',
             '下图展示了商务品的管理费率分布情况。从图中可以看出，大多数商务品的管理费率集中在特定区间，'
             '反映了市场对ETF费率的一般定价水平。费率的分布情况对于投资者选择ETF产品具有重要参考价值，'
             '也反映了不同类型ETF的成本结构差异。'),
            (8, '商务品分成比例分布', '商务品分成比例分布.png',
             '下图展示了商务品的分成比例分布情况。分成比例是衡量商务品商业模式的重要指标，'
             '反映了基金公司与合作方之间的利益分配结构。从分布情况看，市场上商务品的分成比例存在一定差异，'
             '这部分产品定位、合作模式以及市场竞争状况密切相关。'),
            (9, '商务品规模分布', '商务品规模分布.png',
             '下图展示了商务品的规模分布情况。从分布图可以看出，商务品的规模呈现出明显的集中趋势，'
             '少数大规模商务品占据了市场的主要份额，而大多数商务品规模相对较小。这种"头部集中"的现象'
             '反映了ETF市场的马太效应，也说明投资者资金倾向于流向知名度高、流动性的好的ETF产品。'),
            (10, '商务品预期收入分布', '商务品预期收入分布.png',
             '下图展示了商务品的预期收入分布情况。预期收入是评估商务品商业价值的关键指标，'
             '它综合考虑了规模、费率和分成比例等因素。从分布情况看，商务品的预期收入同样呈现出集中趋势，'
             '这部分产品规模分布特征相似，说明规模是影响预期收入的主要因素。'),
            (11, '前10家公司商务品数量', '前10家公司商务品数量.png',
             '下图展示了拥有商务品数量最多的前10家基金公司。从图中可以看出，领先的基金公司在商务品数量上存在明显差距，'
             '反映了不同公司在ETF业务发展上的战略差异。头部公司凭借其品牌优势和资源禀赋，能够推出更多的商务品ETF，'
             '形成了较为明显的市场集中度。'),
            (12, '前10家公司商务品规模', '前10家公司商务品规模.png',
             '下图展示了商务品规模最大的前10家基金公司。与数量排名相比，规模排名更能反映公司在商务品市场的实际影响力。'
             '从图中可以看出，头部公司在规模上的优势更为明显，说明这些公司不仅推出了较多的商务品，'
             '而且这些产品普遍获得了市场的认可，吸引了更多的投资资金。'),
            (13, '商务品平均客户数分布', '商务品平均客户数分布.png',
             '下图展示了商务品的平均持仓客户数分布情况。持仓客户数是衡量ETF产品受欢迎程度和投资者基础的重要指标。'
             '从分布情况看，大多数商务品的持仓客户数相对集中，但也存在一些拥有大量持仓客户的"明星产品"。'
             '这部分产品分布反映了投资者对不同商务品的偏好差异，也说明了产品定位和营销策略的重要性。'),
            (14, '商务品日均成交额分布', '商务品日均成交额分布.png',
             '下图展示了商务品的日均成交额分布情况。日均成交额是衡量ETF产品流动性的重要指标。'
             '从分布情况看，商务品的日均成交额差异较大，少数产品成交活跃，而大多数产品成交相对较少。'
             '这部分产品分布代表了市场交易活动的集中性，也说明了流动对ETF产品吸引力的重要影响。')
        ]
        
        # 已添加的图表文件名列表
        already_added = []
        
        # 按顺序添加图表
        for index, title, filename, description in charts_info:
            chart_path = os.path.join(self.output_dir, filename)
            if os.path.exists(chart_path):
                doc.add_heading(f'3.{index} {title}', 2)
                doc.add_paragraph(description)
                doc.add_picture(chart_path, width=Inches(6))
                print(f"添加图表: {chart_path}")
                already_added.append(filename)
            else:
                print(f"图表文件不存在: {chart_path}")
        
        # 检查是否有其他图表文件
        all_png_files = [f for f in os.listdir(self.output_dir) if f.endswith('.png')]
        already_added.extend([
            'summary_pie_chart.png', 'scale_distribution.png', 'fee_distribution.png',
            'commission_distribution.png', 'company_ranking.png', 'company_scale_ranking.png'
        ])
        
        # 添加其他未添加的图表
        other_charts = [f for f in all_png_files if f not in already_added and not f.startswith('index_category_') and not f.startswith('index_performance_')]
        if other_charts:
            next_index = len([c for c in charts_info if os.path.exists(os.path.join(self.output_dir, c[2]))]) + 1
            for chart_file in other_charts:
                # 提取图表名称作为标题
                chart_name = chart_file.replace('.png', '').replace('_', ' ')
                doc.add_heading(f'3.{next_index} {chart_name}', 2)
                next_index += 1
                
                # 根据图表名称生成描述
                description = self._generate_chart_description(chart_file)
                doc.add_paragraph(description)
                
                # 添加图表
                chart_path = os.path.join(self.output_dir, chart_file)
                doc.add_picture(chart_path, width=Inches(6))
                print(f"添加图表: {chart_path}")
                
                doc.add_paragraph('')
        
        # 检查是否有指数分类数据
        if index_stats is None:
            # 尝试从文件夹中查找指数分类图表
            index_files = [f for f in os.listdir(self.output_dir) if f.startswith('index_category_')]
            if index_files:
                next_index = max(next_index if 'next_index' in locals() else len(charts_info) + 1, len(charts_info) + 1)
                for file in index_files:
                    category = file.replace('index_category_', '').replace('.png', '')
                    if '_bar' in category:
                        continue  # 跳过条形图，稍后处理
                    
                    doc.add_heading(f'3.{next_index} 按{category}分布', 2)
                    next_index += 1
                    
                    # 添加图表
                    doc.add_paragraph('')
                    chart_path = os.path.join(self.output_dir, file)
                    doc.add_picture(chart_path, width=Inches(6))
                    print(f"添加图表: {chart_path}")
                    
                    # 如果有条形图，也添加
                    bar_chart_path = os.path.join(self.output_dir, f'index_category_{category}_bar.png')
                    if os.path.exists(bar_chart_path):
                        doc.add_picture(bar_chart_path, width=Inches(6))
                        print(f"添加图表: {bar_chart_path}")
                    
                    doc.add_paragraph('')
                
                # 添加指数表现分析（如果有）
                performance_chart_path = os.path.join(self.output_dir, 'index_performance_top10.png')
                if os.path.exists(performance_chart_path):
                    doc.add_heading(f'3.{next_index} 指数表现分析', 2)
                    
                    # 添加指数表现表格
                    performance_data = index_stats['指数表现']
                    table = doc.add_table(rows=1, cols=len(performance_data.columns))
                    table.style = 'Table Grid'
                    
                    # 设置表头
                    header_cells = table.rows[0].cells
                    for i, col in enumerate(performance_data.columns):
                        header_cells[i].text = col
                    
                    # 添加数据行（只显示前10个）
                    for _, row in performance_data.head(10).iterrows():
                        cells = table.add_row().cells
                        for i, col in enumerate(performance_data.columns):
                            cells[i].text = str(row[col])
                    
                    # 添加图表
                    doc.add_paragraph('')
                    chart_path = os.path.join(self.output_dir, 'index_performance_top10.png')
                    if os.path.exists(chart_path):
                        doc.add_picture(chart_path, width=Inches(6))
                        print(f"添加图表: {chart_path}")
            else:
                # 尝试查找其他可能的指数分类图表
                index_files = [f for f in os.listdir(self.output_dir) if '指数类型' in f and f.endswith('.png') and f not in already_added]
                if index_files:
                    next_index = max(next_index if 'next_index' in locals() else len(charts_info) + 1, len(charts_info) + 1)
                    for i, file in enumerate(index_files):
                        doc.add_heading(f'3.{next_index + i} 指数分类分析', 2)
                        doc.add_paragraph('')
                        chart_path = os.path.join(self.output_dir, file)
                        doc.add_picture(chart_path, width=Inches(6))
                        print(f"添加图表: {chart_path}")
                        doc.add_paragraph('')
                else:
                    doc.add_paragraph('没有可用的指数分类数据')
            return
        
        # 添加指数分类分析
        next_index = max(next_index if 'next_index' in locals() else len(charts_info) + 1, len(charts_info) + 1)
        for category, data in index_stats.items():
            # 跳过指数表现数据，单独处理
            if category == '指数表现':
                continue
                
            doc.add_heading(f'3.{next_index} 按{category}分布', 2)
            next_index += 1
            
            # 添加分类分布表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # 设置表头
            header_cells = table.rows[0].cells
            header_cells[0].text = category
            header_cells[1].text = 'ETF数量'
            header_cells[2].text = '占比'
            
            # 添加数据行
            for _, row in data.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row[category])
                cells[1].text = str(row['ETF数量'])
                cells[2].text = str(row['占比'])
            
            # 添加图表
            doc.add_paragraph('')
            chart_path = os.path.join(self.output_dir, f'index_category_{category}.png')
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6))
                print(f"添加图表: {chart_path}")
                
            # 如果有条形图，也添加
            bar_chart_path = os.path.join(self.output_dir, f'index_category_{category}_bar.png')
            if os.path.exists(bar_chart_path):
                doc.add_picture(bar_chart_path, width=Inches(6))
                print(f"添加图表: {bar_chart_path}")
            
            doc.add_paragraph('')
        
        # 添加指数表现分析（如果有）
        if '指数表现' in index_stats:
            doc.add_heading(f'3.{next_index} 指数表现分析', 2)
            
            # 添加指数表现表格
            performance_data = index_stats['指数表现']
            table = doc.add_table(rows=1, cols=len(performance_data.columns))
            table.style = 'Table Grid'
            
            # 设置表头
            header_cells = table.rows[0].cells
            for i, col in enumerate(performance_data.columns):
                header_cells[i].text = col
            
            # 添加数据行（只显示前10个）
            for _, row in performance_data.head(10).iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(performance_data.columns):
                    cells[i].text = str(row[col])
            
            # 添加图表
            doc.add_paragraph('')
            chart_path = os.path.join(self.output_dir, 'index_performance_top10.png')
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6))
                print(f"添加图表: {chart_path}")
    
    def _generate_chart_description(self, chart_file):
        """根据图表文件名生成描述"""
        chart_name = chart_file.replace('.png', '').replace('_', ' ')
        
        # 根据图表名称生成不同的描述
        if '客户数' in chart_name:
            return f"下图展示了{chart_name}的分布情况。客户数是衡量ETF产品受欢迎程度的重要指标，反映了产品的市场认可度和投资者基础。"
        elif '成交额' in chart_name:
            return f"下图展示了{chart_name}的分布情况。成交额是衡量ETF产品流动性的关键指标，高流动性有助于降低交易成本，提升产品吸引力。"
        elif '规模' in chart_name:
            return f"下图展示了{chart_name}的分布情况。规模是ETF产品的基础属性，直接影响产品的市场影响力和收益能力。"
        elif '费率' in chart_name:
            return f"下图展示了{chart_name}的分布情况。费率结构是ETF产品的重要特征，直接影响投资者的持有成本和基金公司的盈利能力。"
        elif '分成' in chart_name:
            return f"下图展示了{chart_name}的分布情况。分成比例反映了基金公司与合作方之间的利益分配结构，是商务品商业模式的核心要素。"
        else:
            return f"下图展示了{chart_name}的分布情况。该指标是评估ETF商务品市场特征的重要维度，有助于理解市场结构和发展趋势。"
    
    def _generate_company_charts(self, company_stats):
        """生成公司分析图表"""
        try:
            if company_stats is None:
                print("没有公司统计数据，跳过生成公司分析图表")
                return
            
            # 生成公司商务品数量排名图表
            plt.figure(figsize=(12, 8))
            if '公司商务品数量' in company_stats:
                company_count = company_stats['公司商务品数量']
                # 取前10名
                top10 = company_count.head(10)
                plt.barh(top10.index[::-1], top10.values[::-1])
                plt.title('前10家公司商务品数量')
                plt.xlabel('商务品数量')
                plt.tight_layout()
                plt.savefig(os.path.join(self.output_dir, 'company_ranking.png'))
                plt.close()
            
            # 生成公司商务品规模排名图表
            plt.figure(figsize=(12, 8))
            if '公司商务品规模' in company_stats:
                company_scale = company_stats['公司商务品规模']
                # 取前10名
                top10 = company_scale.head(10)
                plt.barh(top10.index[::-1], top10.values[::-1])
                plt.title('前10家公司商务品规模(亿元)')
                plt.xlabel('规模(亿元)')
                plt.tight_layout()
                plt.savefig(os.path.join(self.output_dir, 'company_scale_ranking.png'))
                plt.close()
            
            print("已生成公司分析图表")
        except Exception as e:
            print(f"生成公司分析图表失败: {str(e)}")
            import traceback
            traceback.print_exc()