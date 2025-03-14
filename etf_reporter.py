import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import glob
from datetime import datetime, timedelta

# 新增版本信息
__version__ = "2.7.5"   
RELEASE_DATE = "2025-03-13"  # 请根据实际发布日期修改

def show_version():
    """显示版本信息"""
    print(f"ETF Reporter Version: {__version__} ({RELEASE_DATE})")
    print("Copyright © 2024 邱超. All rights reserved.")

def preprocess_data():
    # 动态获取最新数据文件
    files = glob.glob('ETF_基础数据合并_*.csv')
    if not files:
        raise FileNotFoundError("未找到ETF基础数据文件")
    
    # 提取最新文件日期
    filename = sorted(files)[-1]
    date_str = filename.split('_')[-1].split('.')[0]
    
    try:
        current_date = datetime.strptime(date_str, "%Y%m%d")
    except ValueError:
        raise ValueError("文件名日期格式应为ETF_基础数据合并_YYYYMMDD.csv")
    
    previous_date = current_date - timedelta(days=7)
    
    # 修正日期格式（保留前导零）
    current_str = current_date.strftime("%m月%d日")  # 02月14日（保留前导零）
    previous_str = previous_date.strftime("%m月%d日")  # 02月07日

    # 读取数据
    result = pd.read_csv(filename, encoding='utf-8-sig')
    
    # 动态商务协议文件
    商务协议_path = f'/Users/admin/Downloads/ETF单产品商务协议{date_str}.xlsx'
    try:
        # 尝试不同的方式读取Excel文件
        商务协议 = pd.read_excel(商务协议_path, engine='openpyxl')
        
        # 详细检查商务协议文件
        print(f"商务协议文件路径: {商务协议_path}")
        print(f"商务协议文件总行数: {len(商务协议)}")
        print(f"商务协议文件列名: {商务协议.columns.tolist()}")
        
        # 检查是否有重复的证券代码
        重复代码 = 商务协议['证券代码'].duplicated().sum()
        print(f"商务协议中重复的证券代码数量: {重复代码}")
        
        # 找出重复的证券代码
        重复代码列表 = 商务协议[商务协议['证券代码'].duplicated(keep=False)]['证券代码'].unique()
        print(f"重复的证券代码列表: {list(重复代码列表)}")
        
        # 显示每个重复代码的详细信息
        print("重复证券代码的详细信息:")
        for code in 重复代码列表:
            重复项 = 商务协议[商务协议['证券代码'] == code]
            print(f"证券代码 {code} 出现 {len(重复项)} 次:")
            for idx, row in 重复项.iterrows():
                print(f"  - 行 {idx+1}: {row['产品名称']} ({row['基金公司简称']})")
        
        # 检查是否有空值
        空值数量 = 商务协议['证券代码'].isna().sum()
        print(f"商务协议中证券代码为空的数量: {空值数量}")
        
        # 打印前几行数据进行检查
        print("商务协议前5行数据:")
        print(商务协议.head(5))
        
        # 确保证券代码为字符串并去除空格
        商务协议['证券代码'] = 商务协议['证券代码'].astype(str).str.strip()
        
        # 再次检查唯一证券代码数量
        唯一代码数量 = len(商务协议['证券代码'].unique())
        print(f"商务协议中唯一证券代码数量: {唯一代码数量}")
        
    except FileNotFoundError:
        raise FileNotFoundError(f"商务协议文件缺失：{商务协议_path}")
    except Exception as e:
        raise Exception(f"读取商务协议文件出错: {str(e)}")

    # 合并商务协议数据前的处理
    result['证券代码'] = result['证券代码'].astype(str).str.strip()
    商务协议['是否商务品'] = '商务'
    
    # 检查合并前后的商务品数量
    商务品代码集合 = set(商务协议['证券代码'].unique())
    基础数据代码集合 = set(result['证券代码'].unique())
    未匹配商务品 = 商务品代码集合 - 基础数据代码集合
    
    if len(未匹配商务品) > 0:
        print(f"警告: 有 {len(未匹配商务品)} 个商务品在基础数据中未找到匹配项")
        print(f"未匹配的商务品代码: {sorted(list(未匹配商务品))}")
    
    # 合并商务协议数据
    result = result.merge(
        商务协议[['证券代码', '是否商务品']],
        on='证券代码',
        how='left'
    ).fillna({'是否商务品': '非商务'})
    
    # 打印合并后的商务品数量（调试信息）
    print(f"合并后的商务品数量: {len(result[result['是否商务品'] == '商务']['证券代码'].unique())}")
    
    # 读取分类数据
    classification_path = '/Users/admin/Downloads/ETF-Index-Classification_20250307.xlsx'
    classification_df = pd.read_excel(classification_path, engine='openpyxl')
    classification_df['跟踪指数代码'] = classification_df['跟踪指数代码'].astype(str).str.strip()

    # 合并分类数据（假设基础数据中有“跟踪指数代码”列）
    result['跟踪指数代码'] = result['跟踪指数代码'].astype(str).str.strip()
    result = result.merge(
        classification_df[['跟踪指数代码', '一级分类', '二级分类', '三级分类']],
        on='跟踪指数代码',
        how='left'
    )

    # 填充空值
    result['基金管理人'] = result['基金管理人'].fillna('未知管理人')
    result['跟踪指数名称'] = result['跟踪指数名称'].fillna('未知指数')

    # 动态生成数值列（关键修复）
    numeric_cols = []
    for metric in ['关注人数', '持仓客户数', '保有金额']:
        # 确保列名格式与数据完全一致
        current_col = f'{metric}（{current_str}）'
        previous_col = f'{metric}（{previous_str}）'
        
        # 验证列是否存在
        if current_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{current_col}'不存在，可用列：\n{available_cols}")
        if previous_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{previous_col}'不存在，可用列：\n{available_cols}")
        
        numeric_cols.extend([current_col, previous_col])
        result[f'{metric}变动'] = result[current_col] - result[previous_col]

    # 数值类型转换
    result[numeric_cols] = result[numeric_cols].fillna(0).astype(int)

    # 生成正确日期范围（MMDD-MMDD）
    date_range = f"{previous_date.strftime('%m%d')}-{current_date.strftime('%m%d')}"
    return result, date_range

def get_manager_short(full_name):
    """基金管理人简称提取"""
    special_mapping = {
        '汇添富基金管理有限公司': '汇添富',
        '易方达基金管理有限公司': '易方达',
        '华夏基金管理有限公司': '华夏',
        '南方基金管理股份有限公司': '南方',
        '嘉实基金管理有限公司': '嘉实'
    }
    if full_name in special_mapping:
        return special_mapping[full_name]

    patterns = ['基金管理有限公司', '基金管理公司', '基金', '管理有限公司', '股份有限公司']
    for pattern in patterns:
        if pattern in full_name:
            return full_name.split(pattern)[0]

    return full_name[:4] if len(full_name) >= 4 else full_name

class ETFReporter:
    def __init__(self, data, date_range):
        self.doc = Document()
        self.data = data
        self.date_range = date_range

        # 解析日期范围（保持前导零）
        start_str, end_str = date_range.split('-')
        self.start_date = f"{start_str[:2]}月{start_str[2:]}日"  # 02月07日
        self.end_date = f"{end_str[:2]}月{end_str[2:]}日"      # 02月14日
        
        self._set_styles()

    def _set_styles(self):
        # 设置正文样式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        font.size = Pt(10.5)
        
        # 设置标题样式，确保所有标题也使用微软雅黑
        for i in range(1, 10):  # 设置标题1-9级
            if f'Heading {i}' in self.doc.styles:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_font = heading_style.font
                heading_font.name = '微软雅黑'
                heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def generate_report(self):
        self._add_cover()
        self._add_business_analysis()
        self._add_combined_analysis('关注人数', '关注人次', {'total': '', 'delta': ''})
        self._add_combined_analysis('持仓客户数', '持仓客户数', {'total': '户', 'delta': '户'})
        self._add_combined_analysis('保有金额', '保有金额', {'total': '亿', 'delta': '百万'})
        # 新增跟踪指数统计
        self._add_tracking_index_stats()
        self.add_classification_analysis()
        self._add_tables()
        
        # 将基金公司商务品收入统计移到最后
        self._add_company_income_stats()
        
        # 添加页码
        self._add_page_numbers()
        
        self.doc.save(f'ETF产品运营周报（{self.date_range}）.docx')
        print(f"周报生成完成：ETF产品运营周报（{self.date_range}）.docx")
        
    def _add_page_numbers(self):
        """为文档添加页码"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        
        # 添加节，以便设置页码
        section = self.doc.sections[0]
        
        # 设置页脚
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        # 设置页脚段落居中
        paragraph.alignment = 1  # 1表示居中对齐
        
        # 添加页码域代码
        run = paragraph.add_run()
        
        # 添加页码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # 设置页码字体为微软雅黑
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(9)  # 设置页码字体大小

    def _add_cover(self):
        self.doc.add_heading('ETF产品运营周报', 0)
        p = self.doc.add_paragraph()
        
        # 新增版本信息
        version_info = f'报告版本：{__version__} ({RELEASE_DATE})\n'
        p.add_run(version_info).font.size = Pt(9)
        p.add_run(version_info).italic = True
        
        date_range_str = f'报告时间范围：2025年{self.start_date}至{self.end_date}\n'
        p.add_run(date_range_str).bold = True
        p.add_run('汇报人：邱超').italic = True

    def _add_combined_analysis(self, metric, name, unit_map):

        if metric == '保有金额':
            unit_map = {'total': '', 'delta': ''}  # 单位由_format_value处理
        total_data = self.data
        current_col = f'{metric}（{self.end_date}）'
        previous_col = f'{metric}（{self.start_date}）'
        
        total = total_data[current_col].sum()
        delta = total_data[f'{metric}变动'].sum()

        biz_data = self.data[self.data['是否商务品'] == '商务']
        biz_total = biz_data[current_col].sum()
        biz_delta = biz_data[f'{metric}变动'].sum()

        self.doc.add_heading(f"{name}分析", level=2)
        p = self.doc.add_paragraph()

        # 总体描述
        total_str = self._format_value(total, metric, is_total=True) + unit_map['total']
        delta_str = self._format_delta(delta, unit_map['delta'], metric)
        main_text = f"截止{self.end_date}，{name}总量为{total_str}，本周变动：{delta_str}。"
        run = p.add_run(main_text)
        run.bold = True

        self._add_top_bottom(p, total_data, metric)

        if not biz_data.empty:
            # 商务部分
            biz_total_str = self._format_value(biz_total, metric, is_total=True) + unit_map['total']
            biz_delta_str = self._format_delta(biz_delta, unit_map['delta'], metric)

            biz_p = self.doc.add_paragraph()
            biz_text = f"其中，商务品{name}总量为{biz_total_str}，本周变动：{biz_delta_str}。"
            biz_run = biz_p.add_run(biz_text)
            biz_run.bold = True

            self._add_top_bottom(biz_p, biz_data, metric, is_business=True)

    def _format_value(self, value, metric, is_total=False):
        """智能数值格式化（修正版本）"""
        if metric == '保有金额':
            if is_total:  # 总量处理
                return f"{value/1e8:,.2f}亿"  # 强制保留两位小数
            # 变动值处理
            if abs(value) >= 1e8:
                return f"{value/1e8:,.2f}亿"
            elif abs(value) >= 1e4:
                return f"{value/1e4:,.0f}万"
            return f"{value:,.0f}元"
        elif metric == '持仓客户数':  # 户数处理
            return f"{int(value):,}"  # 强制转换为整数
        return f"{int(value):,}"  # 其他数值类型也取整

    def _format_delta(self, delta, unit, metric):
        """带符号的变动值格式化（修正版本）"""
        sign = '+' if delta >= 0 else '-'
        abs_value = abs(delta)
        
        # 统一使用_format_value处理数值
        formatted = self._format_value(abs_value, metric)
        return f"{sign}{formatted}"

    def _add_top_bottom(self, paragraph, data, metric, is_business=False):
        top3 = data.nlargest(3, f'{metric}变动')
        bottom3 = data.nsmallest(3, f'{metric}变动')

        increase_str = self._format_products(top3, metric, is_business)
        decrease_str = self._format_products(bottom3, metric, is_business)

        paragraph.add_run(f"增加前三：{increase_str}。")
        paragraph.add_run(f"减少前三：{decrease_str}。")

    def _format_products(self, df, metric, is_business):
        """产品条目格式化"""
        items = []
        for _, row in df.iterrows():
            manager = get_manager_short(row['基金管理人'])
            delta = row[f'{metric}变动']
            delta_str = self._format_delta(delta, '', metric)

            components = [delta_str, manager]
            if not is_business:
                components.append(row['是否商务品'])

            item_str = f"{row[f'证券名称（{self.end_date}）']} ({'、'.join(components)})"
            items.append(item_str)
        return "，".join(items)

    def _add_tracking_index_stats(self):
        """新增跟踪指数统计模块"""
        self.doc.add_heading("跟踪指数统计", level=2)
        
        # 新增：读取THE_BEST_ETF_FINAL.xlsx文件获取指数分级数据
        try:
            best_etf_file = 'THE_BEST_ETF_FINAL.xlsx'
            
            # 确保安装了openpyxl
            try:
                import openpyxl
                print(f"openpyxl版本: {openpyxl.__version__}")
            except ImportError:
                print("警告: openpyxl未安装，尝试使用其他引擎读取Excel")
            
            # 尝试不同的引擎读取Excel文件
            try:
                best_etf_data = pd.read_excel(best_etf_file, sheet_name='优选ETF', engine='openpyxl')
                print(f"成功使用openpyxl引擎读取Excel文件，包含{len(best_etf_data)}行数据")
            except Exception as e1:
                print(f"使用openpyxl引擎读取失败: {str(e1)}")
                try:
                    best_etf_data = pd.read_excel(best_etf_file, sheet_name='优选ETF', engine='xlrd')
                    print(f"成功使用xlrd引擎读取Excel文件，包含{len(best_etf_data)}行数据")
                except Exception as e2:
                    print(f"使用xlrd引擎读取失败: {str(e2)}")
                    # 最后尝试不指定引擎
                    best_etf_data = pd.read_excel(best_etf_file, sheet_name='优选ETF')
                    print(f"成功使用默认引擎读取Excel文件，包含{len(best_etf_data)}行数据")
            
            # 打印列名，帮助调试
            print(f"Excel文件的列名: {best_etf_data.columns.tolist()}")
            
            # 检查'跟踪指数分级'列是否存在
            if '跟踪指数分级' not in best_etf_data.columns:
                print(f"警告: '跟踪指数分级'列不存在，可用列: {best_etf_data.columns.tolist()}")
                raise KeyError("'跟踪指数分级'列不存在")
            
            # 计算各级别指数数量
            total_indices = len(best_etf_data)
            
            # 确保分级数据是字符串格式
            best_etf_data['跟踪指数分级'] = best_etf_data['跟踪指数分级'].astype(str)
            
            # 打印分级数据的唯一值，帮助调试
            print(f"分级数据的唯一值: {best_etf_data['跟踪指数分级'].unique().tolist()}")
            
            absolute_advantage = len(best_etf_data[best_etf_data['跟踪指数分级'].isin(['一级', '二级', '三级', '四级'])])
            relative_advantage = len(best_etf_data[best_etf_data['跟踪指数分级'].isin(['五级', '六级', '七级'])])
            only_business = len(best_etf_data[best_etf_data['跟踪指数分级'] == '八级'])
            no_business = len(best_etf_data[best_etf_data['跟踪指数分级'] == '九级'])
            
            # 计算占比
            absolute_percentage = (absolute_advantage / total_indices) * 100
            relative_percentage = (relative_advantage / total_indices) * 100
            only_business_percentage = (only_business / total_indices) * 100
            no_business_percentage = (no_business / total_indices) * 100
            
            # 添加指数分级统计描述
            p = self.doc.add_paragraph()
            index_stats_text = f"目前全市场共有跟踪指数 {int(total_indices)} 个。其中，我司绝对占优（1-4级）指数有 {int(absolute_advantage)} 个，"
            index_stats_text += f"占比 {absolute_percentage:.1f}%；相对占优（5-7级）指数有 {int(relative_advantage)} 个，"
            index_stats_text += f"占比 {relative_percentage:.1f}%；剩余无优势仅有商务（8级）指数有 {int(only_business)} 个，"
            index_stats_text += f"占比 {only_business_percentage:.1f}%；无商务指数有 {int(no_business)} 个，"
            index_stats_text += f"占比 {no_business_percentage:.1f}%。"
            
            p.add_run(index_stats_text).bold = True
            
        except Exception as e:
            print(f"读取THE_BEST_ETF_FINAL.xlsx文件失败: {str(e)}")
            import traceback
            traceback.print_exc()
            print("跳过指数分级统计部分")
        
        # 按跟踪指数分组计算各项指标
        grouped = self.data.groupby('跟踪指数名称').agg({
            '关注人数变动': 'sum',
            '持仓客户数变动': 'sum',
            '保有金额变动': 'sum'
        }).reset_index()

        # 定义统计指标和对应的中文名称
        metrics = [
            ('关注人数变动', '新增关注数', '取消关注'),
            ('持仓客户数变动', '持仓客户数增长', '持仓客户数减少'),
            ('保有金额变动', '保有金额增长', '保有金额减少')
        ]

        # 为每个指标生成统计结果
        for metric, pos_name, neg_name in metrics:
            # 获取前三和后三
            top3 = grouped.nlargest(3, metric)
            bottom3 = grouped.nsmallest(3, metric)
            
            # 格式化结果
            top_str = self._format_index_changes(top3, metric)
            bottom_str = self._format_index_changes(bottom3, metric)
            
            # 添加段落
            p = self.doc.add_paragraph()
            p.add_run(f"{pos_name}前三：{top_str}。{neg_name}前三：{bottom_str}。")

    def _format_index_changes(self, df, metric):
        """格式化指数变动信息（修正版本）"""
        items = []
        for _, row in df.iterrows():
            value = row[metric]
            sign = '+' if value >= 0 else '-'
            
            # 根据指标类型确定格式化方式
            metric_type = '保有金额' if '金额' in metric else metric.replace('变动', '')
            formatted = self._format_value(abs(value), metric_type)
            
            items.append(f"{row['跟踪指数名称']}（{sign}{formatted}）")
        return "，".join(items)

    def _add_tables(self):
        for metric in ['关注人数', '持仓客户数', '保有金额']:
            top20 = self.data.nlargest(20, f'{metric}变动')
            self._add_table(f'{metric}增长Top20', top20, metric)

            bottom20 = self.data.nsmallest(20, f'{metric}变动')
            self._add_table(f'{metric}减少Top20', bottom20, metric)

    def _add_table(self, title, data, metric):
        self.doc.add_heading(title, level=2)
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading Accent 1'

        headers = ['证券代码', '产品名称', '管理人', '当前值', '上周值', '变动值', '商务属性']
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # 初始化序号计数器 (修复点)
        desc_counter = 1  # 在循环外部初始化

        for _, row in data.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['证券代码'])
            cells[1].text = row[f'证券名称（{self.end_date}）']
            cells[2].text = get_manager_short(row['基金管理人'])
            cells[3].text = self._format_value(row[f'{metric}（{self.end_date}）'], metric)
            cells[4].text = self._format_value(row[f'{metric}（{self.start_date}）'], metric)
            cells[5].text = self._format_delta(row[f'{metric}变动'], '', metric)
            cells[6].text = row['是否商务品']

            # 针对非商务品，生成描述信息
            if row['是否商务品'] == '非商务' and title == "关注人数增长Top20":
                # 处理数值
                month_avg = row['月成交额[交易日期]最新收盘日[单位]百万元']
                rounded_month_avg = int(round(month_avg))
                management_fee = row['管理费率[单位]%']
                metric_change = int(row[f'{metric}变动'])

                # 添加描述段落
                para = self.doc.add_paragraph()
                para.add_run(f"{desc_counter}. ").bold = True
                para.add_run(f"{get_manager_short(row['基金管理人'])}基金的{row[f'证券名称（{self.end_date}）']}（代码{row['证券代码']}），月日均交易额{rounded_month_avg}，本周新增{metric_change}关注数，管理费率{management_fee}%，非商务品。以下是同跟踪指数产品列表：")
                
                self._add_alternative_etfs(row['跟踪指数代码'], metric)
                desc_counter += 1

    def _add_alternative_etfs(self, tracking_index_code, metric):
        # 获取同跟踪指数的所有ETF产品
        alternative_etfs = self.data[self.data['跟踪指数代码'] == tracking_index_code]
        
        # 去重并排序
        alternative_etfs = alternative_etfs.drop_duplicates(subset=['证券代码'])
        alternative_etfs = alternative_etfs.sort_values(by='月成交额[交易日期]最新收盘日[单位]百万元', ascending=False)
        
        # 只展示前十个
        alternative_etfs = alternative_etfs.head(10)
        
        # 生成表格
        table = self.doc.add_table(rows=1, cols=10)
        table.style = 'Light Shading Accent 1'
        
        headers = ['产品代码', '产品名称', '管理人', '关注数当前值', '关注数变动值', '规模', '管理费率', '月日均交易额', '商务属性']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for _, row in alternative_etfs.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['证券代码'])
            cells[1].text = row[f'证券名称（{self.end_date}）']
            cells[2].text = get_manager_short(row['基金管理人'])
            cells[3].text = self._format_value(row[f'{metric}（{self.end_date}）'], metric)
            cells[4].text = self._format_delta(row[f'{metric}变动'], '', metric)
            cells[5].text = self._format_value(row['基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'], '规模')
            cells[6].text = str(row['管理费率[单位]%'])
            cells[7].text = self._format_value(row['月成交额[交易日期]最新收盘日[单位]百万元'], '月日均交易额')            
            cells[8].text = row['是否商务品']

    def add_classification_analysis(self):
        # 按分类统计
        for level in ['一级分类', '二级分类', '三级分类']:
            grouped = self.data.groupby(level).agg({
                f'关注人数（{self.end_date}）': 'sum',
                '关注人数变动': 'sum',
                f'持仓客户数（{self.end_date}）': 'sum',
                '持仓客户数变动': 'sum',
                f'保有金额（{self.end_date}）': 'sum',
                '保有金额变动': 'sum'
            }).reset_index()

            # 按新增关注数排序
            grouped = grouped.sort_values(by='关注人数变动', ascending=False)

            # 如果是三级分类，只展示关注数增长前 20 和减少前 10
            if level == '三级分类':
                top20 = grouped.nlargest(20, '关注人数变动')
                bottom10 = grouped.nsmallest(10, '关注人数变动')
                grouped = pd.concat([top20, bottom10])

            # 添加到报告
            self.doc.add_heading(f"{level}统计", level=2)
            table = self.doc.add_table(rows=1, cols=7)
            table.style = 'Light Shading Accent 1'

            headers = [level, '总关注数', '本周关注数变化', '总持仓数', '本周持仓数变化', '总持仓市值', '本周持仓市值变化']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for _, row in grouped.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row[level])
                cells[1].text = str(int(row[f'关注人数（{self.end_date}）']))  # 总关注数
                cells[2].text = str(int(row['关注人数变动']))  # 本周关注数变化（取整数）
                cells[3].text = str(int(row[f'持仓客户数（{self.end_date}）']))  # 总持仓数
                cells[4].text = str(int(row['持仓客户数变动']))  # 本周持仓数变化（取整数）
                cells[5].text = self._format_value(row[f'保有金额（{self.end_date}）'], '保有金额')  # 总持仓市值（智能格式化）
                cells[6].text = self._format_value(row['保有金额变动'], '保有金额')  # 本周持仓市值变化（智能格式化）

# 新增商务品分析方法
    def _add_business_analysis(self):
        """添加商务品分析"""
        self.doc.add_heading("商务品分析", level=2)
        
        # 获取商务品数据
        biz_data = self.data[self.data['是否商务品'] == '商务']
        total_count = len(biz_data['证券代码'].unique())
        
        # 计算商务品在我司的合计规模（使用持仓金额而非基金总规模）
        total_scale = biz_data[f'保有金额（{self.end_date}）'].sum() / 1e8  # 转换为亿元
        
        # 计算所有ETF在我司的总规模
        all_etf_scale = self.data[f'保有金额（{self.end_date}）'].sum() / 1e8  # 转换为亿元
        
        # 计算占比
        percentage = (total_scale / all_etf_scale) * 100 if all_etf_scale > 0 else 0
        
        # 计算预期管理费收入（按照40%分成比例）
        total_expected_income = 0
        
        # 为每个商务品计算预期收入并汇总
        for _, row in biz_data.iterrows():
            # 使用我司持仓规模（保有金额）- 单位是元，需要转换为亿元
            fund_scale = row[f'保有金额（{self.end_date}）'] / 1e8  # 转换为亿元
            management_fee_rate = row['管理费率[单位]%'] / 100  # 转换为小数
            # 计算年化管理费收入：规模(亿元) * 费率 * 40% * 100(转为万元)
            expected_income = fund_scale * management_fee_rate * 0.35 * 10000  # 转换为万元
            total_expected_income += expected_income
                       
        # 计算每个基金公司的预期收入
        # 修改基金公司预期收入计算方式（原代码572行附近）
        company_stats = biz_data.groupby('基金管理人').agg({
            f'保有金额（{self.end_date}）': 'sum',
            f'关注人数（{self.end_date}）': 'sum',
            f'持仓客户数（{self.end_date}）': 'sum'
        }).reset_index()

        # 转换保有金额为亿元
        company_stats[f'保有金额（{self.end_date}）'] = company_stats[f'保有金额（{self.end_date}）'] / 1e8

        # 新增：逐个产品计算后汇总
        product_income = biz_data[f'保有金额（{self.end_date}）'] * biz_data['管理费率[单位]%'] * 0.35 / 1e6  # 直接万元转换
        company_income = biz_data.groupby('基金管理人')[f'保有金额（{self.end_date}）'].apply(
            lambda x: (x * biz_data.loc[x.index, '管理费率[单位]%'] * 0.35 / 1e6).sum()
        ).reset_index(name='预期收入')

        company_stats = company_stats.merge(company_income, on='基金管理人')        
        # 按我司持仓规模排序并获取前三
        top3_companies = company_stats.nlargest(3, f'保有金额（{self.end_date}）')
        
        # 尝试获取全市场ETF数量 - 修改这部分代码以正确获取证券代码数量
        all_market_etf_count = 0
        market_percentage = 0
        
        try:
            # 直接使用指定的ETF数据文件路径
            etf_data_file = '/Users/admin/Downloads/ETF_DATA_20250307.xlsx'
            
            # 确保导入os模块
            import os
            
            if os.path.exists(etf_data_file):
                # 使用pandas直接读取Excel文件
                try:
                    # 读取"万得"sheet
                    etf_df = pd.read_excel(etf_data_file, sheet_name='万得', engine='openpyxl')
                    print(f"成功读取'万得'sheet，包含{len(etf_df)}行数据")
                    
                    # 打印列名，帮助调试
                    print(f"'万得'sheet的列名: {etf_df.columns.tolist()}")
                    
                    # 直接使用"证券代码"列计算ETF数量
                    if '证券代码' in etf_df.columns:
                        # 清理证券代码列
                        valid_codes = etf_df['证券代码'].dropna()
                        valid_codes = valid_codes.astype(str).str.strip()
                        valid_codes = valid_codes[valid_codes != '']
                        valid_codes = valid_codes[valid_codes != 'nan']
                        
                        # 计算唯一证券代码数量，并减去1（去除最后一行数据来源）
                        all_market_etf_count = len(valid_codes.unique()) - 1
                        print(f"从'万得'sheet的'证券代码'列提取到的唯一ETF代码数量(减去数据来源行): {all_market_etf_count}")
                        
                        # 打印部分代码用于验证
                        print(f"部分ETF代码示例: {valid_codes.unique()[:10].tolist()}")
                    else:
                        print("在'万得'sheet中未找到'证券代码'列")
                        all_market_etf_count = 1097  # 使用已知的正确值
                
                except Exception as e:
                    print(f"读取Excel文件失败: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    all_market_etf_count = 1097  # 如果读取失败，使用已知的正确值
                
                # 计算商务品占全市场ETF的比例
                market_percentage = (total_count / all_market_etf_count) * 100
                print(f"商务品占比: {total_count}/{all_market_etf_count} = {market_percentage:.1f}%")
            else:
                print(f"警告: ETF数据文件不存在: {etf_data_file}")
                all_market_etf_count = 1097  # 如果文件不存在，使用已知的正确值
                market_percentage = (total_count / all_market_etf_count) * 100
        except Exception as e:
            print(f"获取全市场ETF数量失败: {str(e)}")
            import traceback
            traceback.print_exc()
            all_market_etf_count = 1097  # 如果出现异常，使用已知的正确值
            market_percentage = (total_count / all_market_etf_count) * 100
        
        # 生成描述文本
        p = self.doc.add_paragraph()
        
        # 新的描述文本格式
        summary_text = f"目前全市场共有 {int(all_market_etf_count)} 只ETF，其中我司商务品有 {int(total_count)} 只，"
        summary_text += f"占比 {market_percentage:.1f}%。我司保有的ETF规模有 {all_etf_scale:.1f} 亿，"
        summary_text += f"其中商务品规模 {total_scale:.1f} 亿，占比 {percentage:.1f}%。"
        summary_text += f"持有所有商务品一年，预估管理费收入为 {int(total_expected_income)} 万元。"
        
        p.add_run(summary_text).bold = True
        
        # 在生成company_text前添加验证
        print("\n基金公司规模验证:")
        print(top3_companies[[f'保有金额（{self.end_date}）', '基金管理人']])

        # 生成前三基金公司描述
        if not top3_companies.empty:
            company_text = "目前，我司商务品规模前三的基金公司是："
            
            for i, (_, row) in enumerate(top3_companies.iterrows()):
                company_name = get_manager_short(row['基金管理人'])
                scale = row[f'保有金额（{self.end_date}）']  # 已转换为亿元
                attention = int(row[f'关注人数（{self.end_date}）'])
                holders = int(row[f'持仓客户数（{self.end_date}）'])
                expected_income = row['预期收入']
                
                company_text += f"{company_name}（我司持仓 {scale:.1f}亿元，关注 {attention:,}人，持仓 {holders:,}人，预估一年收入 {expected_income:.0f} 万）"
                
                if i < len(top3_companies) - 1:
                    company_text += "，"
                else:
                    company_text += "。"
            
            p.add_run(company_text)
            
        # 删除重复的描述和收入统计表部分
        # ...

    def _add_company_income_stats(self):
        """添加基金公司商务品收入统计（作为单独章节放在最后）"""
        self.doc.add_heading("基金公司ETF商务品管理费预计收入统计", level=2)
        
        # 获取商务品数据
        biz_data = self.data[self.data['是否商务品'] == '商务']
        
        # 获取全量数据按公司分组
        all_company = self.data.groupby('基金管理人').agg({
            f'保有金额（{self.end_date}）': 'sum',
            '证券代码': 'nunique'
        }).rename(columns={'证券代码': '所有产品数量'})

        # 商务品数据按公司分组
        biz_company = biz_data.groupby('基金管理人').agg({
            f'保有金额（{self.end_date}）': 'sum',
            '证券代码': 'nunique',
            '管理费率[单位]%': 'mean'
        }).rename(columns={
            '证券代码': '商务产品数量',
            f'保有金额（{self.end_date}）': '商务规模'
        })

        # 合并数据
        merged_stats = all_company.merge(biz_company, on='基金管理人', how='left').fillna(0)
        
        # 计算各项收入（单位：万元）
        for idx, row in merged_stats.iterrows():
            # 获取该基金公司的所有产品
            company_products = self.data[self.data['基金管理人'] == idx]
            # 获取该基金公司的商务品
            company_biz = biz_data[biz_data['基金管理人'] == idx]
            
            # 计算总收入
            total_income = 0
            for _, prod in company_products.iterrows():
                scale = prod[f'保有金额（{self.end_date}）'] / 1e8  # 转换为亿元
                fee_rate = prod['管理费率[单位]%'] / 100  # 转换为小数
                income = scale * fee_rate * 0.35 * 10000  # 转换为万元
                total_income += income
            
            # 计算商务品收入
            biz_income = 0
            for _, prod in company_biz.iterrows():
                scale = prod[f'保有金额（{self.end_date}）'] / 1e8  # 转换为亿元
                fee_rate = prod['管理费率[单位]%'] / 100  # 转换为小数
                income = scale * fee_rate * 0.35 * 10000  # 转换为万元
                biz_income += income
            
            merged_stats.at[idx, '所有保有预计收入'] = total_income
            merged_stats.at[idx, '商务品预计收入'] = biz_income
            merged_stats.at[idx, '非商务品预计收入'] = total_income - biz_income
        
        # 格式化数值
        merged_stats = merged_stats.reset_index()
        merged_stats = merged_stats.sort_values('商务品预计收入', ascending=False)
        merged_stats['所有保有规模'] = merged_stats[f'保有金额（{self.end_date}）'] / 1e8  # 转换为亿元
        merged_stats['商务规模'] = merged_stats['商务规模'] / 1e8  # 转换为亿元
        
        # 新增：过滤掉规模小于1000万(0.1亿)的基金公司
        merged_stats = merged_stats[merged_stats['所有保有规模'] >= 0.1]
        
        # 生成表格
        table = self.doc.add_table(rows=1, cols=8)
        table.style = 'Light Shading Accent 1'
        
        # 设置表头
        headers = ['基金公司', '所有规模(亿)', '商务规模(亿)', '所有产品数', '商务产品数', 
                 '总预计收入(万)', '商务收入(万)', '非商务收入(万)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        # 填充数据
        for _, row in merged_stats.iterrows():
            cells = table.add_row().cells
            cells[0].text = get_manager_short(row['基金管理人'])
            cells[1].text = f"{row['所有保有规模']:,.2f}"
            cells[2].text = f"{row['商务规模']:,.2f}"
            cells[3].text = f"{int(row['所有产品数量'])}"
            cells[4].text = f"{int(row['商务产品数量'])}"
            cells[5].text = f"{row['所有保有预计收入']:,.0f}"
            cells[6].text = f"{row['商务品预计收入']:,.0f}"
            cells[7].text = f"{row['非商务品预计收入']:,.0f}"
        
        # 添加汇总行
        total_row = table.add_row().cells
        total_row[0].text = "总计"
        total_row[0].paragraphs[0].runs[0].font.bold = True
        
        # 计算各列汇总值
        total_row[1].text = f"{merged_stats['所有保有规模'].sum():,.2f}"
        total_row[2].text = f"{merged_stats['商务规模'].sum():,.2f}"
        total_row[3].text = f"{int(merged_stats['所有产品数量'].sum())}"
        total_row[4].text = f"{int(merged_stats['商务产品数量'].sum())}"
        total_row[5].text = f"{merged_stats['所有保有预计收入'].sum():,.0f}"
        total_row[6].text = f"{merged_stats['商务品预计收入'].sum():,.0f}"
        total_row[7].text = f"{merged_stats['非商务品预计收入'].sum():,.0f}"
        
        # 设置汇总行底色
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        for cell in total_row:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

if __name__ == "__main__":
    try:
        df, date_range = preprocess_data()
        reporter = ETFReporter(df, date_range)
        reporter.generate_report()
    except Exception as e:
        print(f"程序运行出错：{str(e)}")
        print("请检查：")
        print("1. 数据文件列名是否形如'关注人数（02月14日）'")
        print("2. 文件命名是否符合ETF_基础数据合并_YYYYMMDD.csv格式")
        print("3. 商务协议文件是否存在且命名正确")
        print("4. 日期是否跨年/跨月（程序自动处理日期范围）")