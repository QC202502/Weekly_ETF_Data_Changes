"""
ETF业务周报生成工具

版本历史：
- v2.0.0 (2025-02-14)
    * 新增商务品分析功能
    * 新增跟踪指数统计
    * 优化数值格式化显示
    
- v1.0.0 (2025-01-01)
    * 首次发布
    * 基础报告生成功能
"""

__version__ = '2.0.0'

# 在文件开头添加版本配置
VERSION_CONFIG = {
    '2.0.0': {
        'features': [
            '支持商务品分析',
            '支持跟踪指数统计',
            '智能数值格式化',
            '支持日期范围自动处理',
            '版本号管理'
        ],
        'date_format': '%m月%d日',
        'supported_metrics': ['关注人数', '持仓客户数', '保有金额']
    }
}

# 添加配置文件
CONFIG = {
    'paths': {
        'output_dir': str(Path.home() / 'Downloads'),  # 使用用户主目录
        'business_file_pattern': 'ETF单产品商务协议{}.xlsx',
        'data_file_pattern': 'ETF_基础数据合并_*.csv',
        'report_pattern': 'ETF产品运营周报（{}）.docx'
    },
    'metrics': {
        '关注人数': {'unit': '', 'display_name': '关注人次'},
        '持仓客户数': {'unit': '户', 'display_name': '持仓客户数'},
        '保有金额': {'unit': '亿', 'display_name': '保有金额'}
    }
}

def check_version():
    """检查当前程序版本"""
    return __version__

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import glob
from datetime import datetime, timedelta
import logging
from pathlib import Path

def setup_logging():
    """配置日志"""
    logging.basicConfig(
        filename='etf_report.log',
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    return logging.getLogger(__name__)

logger = setup_logging()

def preprocess_data():
    logger.info("开始数据预处理")
    # 动态获取最新数据文件
    files = glob.glob('ETF_基础数据合并_*.csv')
    if not files:
        raise FileNotFoundError("未找到ETF基础数据文件")
    
    # 提取最新文件日期
    filename = sorted(files)[-1]
    date_str = filename.split('_')[-1].split('.')[0]
    
    try:
        current_date = datetime.strptime(date_str, "%Y%m%d")
    except ValueError:
        raise ValueError("文件名日期格式应为ETF_基础数据合并_YYYYMMDD.csv")
    
    previous_date = current_date - timedelta(days=7)
    
    # 修正日期格式（保留前导零）
    current_str = current_date.strftime("%m月%d日")  # 02月14日（保留前导零）
    previous_str = previous_date.strftime("%m月%d日")  # 02月07日

    # 读取数据
    result = pd.read_csv(filename, encoding='utf-8-sig')
    
    # 动态商务协议文件
    商务协议_path = Path(CONFIG['paths']['output_dir']) / CONFIG['paths']['business_file_pattern'].format(date_str)
    try:
        商务协议 = pd.read_excel(商务协议_path, engine='openpyxl')
    except FileNotFoundError:
        raise FileNotFoundError(f"商务协议文件缺失：{商务协议_path}")

    # 合并商务协议数据
    result['证券代码'] = result['证券代码'].astype(str).str.strip()
    商务协议['证券代码'] = 商务协议['证券代码'].astype(str).str.strip()
    商务协议['是否商务品'] = '商务'
    result = result.merge(
        商务协议[['证券代码', '是否商务品']],
        on='证券代码',
        how='left'
    ).fillna({'是否商务品': '非商务'})

    # 填充空值
    result['基金管理人'] = result['基金管理人'].fillna('未知管理人')
    result['跟踪指数名称'] = result['跟踪指数名称'].fillna('未知指数')

    # 动态生成数值列（关键修复）
    numeric_cols = []
    for metric in ['关注人数', '持仓客户数', '保有金额']:
        # 确保列名格式与数据完全一致
        current_col = f'{metric}（{current_str}）'
        previous_col = f'{metric}（{previous_str}）'
        
        # 验证列是否存在
        if current_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{current_col}'不存在，可用列：\n{available_cols}")
        if previous_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{previous_col}'不存在，可用列：\n{available_cols}")
        
        numeric_cols.extend([current_col, previous_col])
        result[f'{metric}变动'] = result[current_col] - result[previous_col]

    # 数值类型转换
    result[numeric_cols] = result[numeric_cols].fillna(0).astype(int)

    # 生成正确日期范围（MMDD-MMDD）
    date_range = f"{previous_date.strftime('%m%d')}-{current_date.strftime('%m%d')}"
    logger.info(f"数据预处理完成，处理日期范围：{date_range}")
    return result, date_range

def get_manager_short(full_name):
    """基金管理人简称提取"""
    special_mapping = {
        '汇添富基金管理有限公司': '汇添富',
        '易方达基金管理有限公司': '易方达',
        '华夏基金管理有限公司': '华夏',
        '南方基金管理股份有限公司': '南方',
        '嘉实基金管理有限公司': '嘉实'
    }
    if full_name in special_mapping:
        return special_mapping[full_name]

    patterns = ['基金管理有限公司', '基金管理公司', '基金', '管理有限公司', '股份有限公司']
    for pattern in patterns:
        if pattern in full_name:
            return full_name.split(pattern)[0]

    return full_name[:4] if len(full_name) >= 4 else full_name

class ETFReporter:
    def __init__(self, data, date_range):
        self.version = __version__
        self.config = VERSION_CONFIG[self.version]
        self.doc = Document()
        self.data = data
        self.date_range = date_range
        self.output_dir = Path(CONFIG['paths']['output_dir'])
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # 解析日期范围（保持前导零）
        start_str, end_str = date_range.split('-')
        self.start_date = f"{start_str[:2]}月{start_str[2:]}日"  # 02月07日
        self.end_date = f"{end_str[:2]}月{end_str[2:]}日"      # 02月14日
        
        self._set_styles()
        
        # 缓存常用数据
        self._cache = {}

    def _set_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        font.size = Pt(10.5)

    def get_features(self):
        """获取当前版本支持的功能列表"""
        return self.config['features']

    def generate_report(self):
        try:
            logger.info(f"开始生成报告: {self.date_range}")
            print(f"正在使用 ETF业务周报 {self.version} 生成报告")
            print("支持的功能：")
            for feature in self.get_features():
                print(f"- {feature}")
            
            self._add_cover()
            self._add_combined_analysis('关注人数', '关注人次', {'total': '', 'delta': ''})
            self._add_combined_analysis('持仓客户数', '持仓客户数', {'total': '户', 'delta': '户'})
            self._add_combined_analysis('保有金额', '保有金额', {'total': '亿', 'delta': '百万'})
            # 新增跟踪指数统计
            self._add_tracking_index_stats()
            self._add_tables()
            output_path = self.output_dir / f'ETF产品运营周报（{self.date_range}）.docx'
            self.doc.save(output_path)
            print(f"周报生成完成：{output_path}")
            logger.info(f"报告生成完成: {self.date_range}")
        except Exception as e:
            logger.error(f"报告生成失败: {str(e)}")
            raise

    def _add_cover(self):
        self.doc.add_heading('ETF产品运营周报', 0)
        p = self.doc.add_paragraph()
        date_range_str = f'报告时间范围：2025年{self.start_date}至{self.end_date}\n'
        p.add_run(date_range_str).bold = True
        p.add_run('汇报人：邱超').italic = True

    def _add_combined_analysis(self, metric, name, unit_map):
        # 使用缓存优化计算
        def compute_total():
            return self.data[f'{metric}（{self.end_date}）'].sum()
        
        total = self._get_cached_data(f'{metric}_total', compute_total)
        delta = self.data[f'{metric}变动'].sum()

        biz_data = self.data[self.data['是否商务品'] == '商务']
        biz_total = biz_data[f'{metric}（{self.end_date}）'].sum()
        biz_delta = biz_data[f'{metric}变动'].sum()

        self.doc.add_heading(f"{name}分析", level=2)
        p = self.doc.add_paragraph()

        # 总体描述
        total_str = self._format_value(total, metric, is_total=True) + unit_map['total']
        delta_str = self._format_delta(delta, unit_map['delta'], metric)
        main_text = f"截止{self.end_date}，{name}总量为{total_str}，本周变动：{delta_str}。"
        run = p.add_run(main_text)
        run.bold = True

        self._add_top_bottom(p, self.data, metric)

        if not biz_data.empty:
            # 商务部分
            biz_total_str = self._format_value(biz_total, metric, is_total=True) + unit_map['total']
            biz_delta_str = self._format_delta(biz_delta, unit_map['delta'], metric)

            biz_p = self.doc.add_paragraph()
            biz_text = f"其中，商务品{name}总量为{biz_total_str}，本周变动：{biz_delta_str}。"
            biz_run = biz_p.add_run(biz_text)
            biz_run.bold = True

            self._add_top_bottom(biz_p, biz_data, metric, is_business=True)

    def _format_value(self, value, metric, is_total=False):
        """智能数值格式化"""
        if metric == '保有金额':
            if is_total:  # 总量处理
                return f"{value/1e8:,.2f}亿"
            # 变动值处理
            if abs(value) >= 1e8:
                return f"{value/1e8:,.2f}亿"
            elif abs(value) >= 1e4:
                return f"{value/1e4:,.0f}万"
            return f"{value:,.0f}元"
        elif metric == '持仓客户数':  # 户数处理
            return f"{int(value):,}"  # 强制转换为整数
        return f"{int(value):,}"  # 其他数值类型也取整

    def _format_delta(self, delta, unit, metric):
        """带符号的变动值格式化"""
        sign = '+' if delta >= 0 else '-'
        abs_value = abs(delta)
        formatted = self._format_value(abs_value, metric)
        return f"{sign}{formatted}"

    def _add_top_bottom(self, paragraph, data, metric, is_business=False):
        top3 = data.nlargest(3, f'{metric}变动')
        bottom3 = data.nsmallest(3, f'{metric}变动')

        increase_str = self._format_products(top3, metric, is_business)
        decrease_str = self._format_products(bottom3, metric, is_business)

        paragraph.add_run(f"增加前三：{increase_str}。")
        paragraph.add_run(f"减少前三：{decrease_str}。")

    def _format_products(self, df, metric, is_business):
        """产品条目格式化"""
        items = []
        for _, row in df.iterrows():
            manager = get_manager_short(row['基金管理人'])
            delta = row[f'{metric}变动']
            delta_str = self._format_delta(delta, '', metric)

            components = [delta_str, manager]
            if not is_business:
                components.append(row['是否商务品'])

            item_str = f"{row[f'证券名称（{self.end_date}）']} ({'、'.join(components)})"
            items.append(item_str)
        return "，".join(items)

    def _add_tracking_index_stats(self):
        """新增跟踪指数统计模块"""
        self.doc.add_heading("跟踪指数统计", level=2)
        
        # 按跟踪指数分组计算各项指标
        grouped = self.data.groupby('跟踪指数名称').agg({
            '关注人数变动': 'sum',
            '持仓客户数变动': 'sum',
            '保有金额变动': 'sum'
        }).reset_index()

        # 定义统计指标和对应的中文名称
        metrics = [
            ('关注人数变动', '新增关注数', '取消关注'),
            ('持仓客户数变动', '持仓客户数增长', '持仓客户数减少'),
            ('保有金额变动', '保有金额增长', '保有金额减少')
        ]

        # 为每个指标生成统计结果
        for metric, pos_name, neg_name in metrics:
            # 获取前三和后三
            top3 = grouped.nlargest(3, metric)
            bottom3 = grouped.nsmallest(3, metric)
            
            # 格式化结果
            top_str = self._format_index_changes(top3, metric)
            bottom_str = self._format_index_changes(bottom3, metric)
            
            # 添加段落
            p = self.doc.add_paragraph()
            p.add_run(f"{pos_name}前三：{top_str}。{neg_name}前三：{bottom_str}。")

    def _format_index_changes(self, df, metric):
        """格式化指数变动信息（修正版本）"""
        items = []
        for _, row in df.iterrows():
            value = row[metric]
            sign = '+' if value >= 0 else '-'
            
            # 根据指标类型确定格式化方式
            metric_type = '保有金额' if '金额' in metric else metric.replace('变动', '')
            formatted = self._format_value(abs(value), metric_type)
            
            items.append(f"{row['跟踪指数名称']}（{sign}{formatted}）")
        return "，".join(items)

    def _add_tables(self):
        for metric in ['关注人数', '持仓客户数', '保有金额']:
            top20 = self.data.nlargest(20, f'{metric}变动')
            self._add_table(f'{metric}增长Top20', top20, metric)

            bottom20 = self.data.nsmallest(20, f'{metric}变动')
            self._add_table(f'{metric}减少Top20', bottom20, metric)

    def _add_table(self, title, data, metric):
        self.doc.add_heading(title, level=2)
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading Accent 1'

        headers = ['证券代码', '产品名称', '管理人', '当前值', '上周值', '变动值', '商务属性']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for _, row in data.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['证券代码'])
            cells[1].text = row[f'证券名称（{self.end_date}）']
            cells[2].text = get_manager_short(row['基金管理人'])
            cells[3].text = self._format_value(row[f'{metric}（{self.end_date}）'], metric)
            cells[4].text = self._format_value(row[f'{metric}（{self.start_date}）'], metric)
            cells[5].text = self._format_delta(row[f'{metric}变动'], '', metric)
            cells[6].text = row['是否商务品']

    def _get_cached_data(self, key, compute_func):
        """使用缓存优化重复计算"""
        if key not in self._cache:
            self._cache[key] = compute_func()
        return self._cache[key]

def validate_data(df):
    """数据验证"""
    required_cols = ['证券代码', '证券名称', '基金管理人', '跟踪指数名称']
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(f"缺少必要列：{', '.join(missing_cols)}")
    
    # 数据类型检查
    if not df['证券代码'].dtype.kind in 'OS':  # Object或String类型
        raise TypeError("证券代码列应为字符串类型")

if __name__ == "__main__":
    try:
        print(f"ETF业务周报 版本 {check_version()}")
        df, date_range = preprocess_data()
        validate_data(df)
        reporter = ETFReporter(df, date_range)
        reporter.generate_report()
    except Exception as e:
        print(f"程序运行出错：{str(e)}")
        print("请检查：")
        print("1. 数据文件列名是否形如'关注人数（02月14日）'")
        print("2. 文件命名是否符合ETF_基础数据合并_YYYYMMDD.csv格式")
        print("3. 商务协议文件是否存在且命名正确")
        print("4. 日期是否跨年/跨月（程序自动处理日期范围）")
        logger.error(f"程序运行出错：{str(e)}")