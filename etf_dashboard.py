import pandas as pd
import glob
import os
import json
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from docx import Document

# 新增版本信息
__version__ = "2.8.0"   
RELEASE_DATE = "2025-03-14"  # 请根据实际发布日期修改

# 导入项目中的其他模块
try:
    from etf_reporter import ETFReporter, get_manager_short
    from business_etf_analyzer import BusinessETFAnalyzer
except ImportError:
    # 如果无法导入，定义一个简单的函数
    def get_manager_short(full_name):
        """基金管理人简称提取"""
        special_mapping = {
            '汇添富基金管理有限公司': '汇添富',
            '易方达基金管理有限公司': '易方达',
            '华夏基金管理有限公司': '华夏',
            '南方基金管理股份有限公司': '南方',
            '嘉实基金管理有限公司': '嘉实'
        }
        if full_name in special_mapping:
            return special_mapping[full_name]

        patterns = ['基金管理有限公司', '基金管理公司', '基金', '管理有限公司', '股份有限公司']
        for pattern in patterns:
            if pattern in full_name:
                return full_name.split(pattern)[0]

        return full_name[:4] if len(full_name) >= 4 else full_name

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 限制上传文件大小为50MB

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 全局变量存储数据
etf_data = None
business_etfs = set()  # 存储商务品ETF代码
current_date_str = ""
previous_date_str = ""
date_range = ""

def load_latest_data():
    """加载最新的ETF数据和商务品信息"""
    global etf_data, business_etfs, current_date_str, previous_date_str, date_range
    
    print("开始加载ETF数据...")
    
    # 动态获取最新数据文件
    files = glob.glob('ETF_基础数据合并_*.csv')
    if not files:
        # 尝试在其他目录查找
        alt_paths = [
            '/Users/admin/Downloads/ETF_基础数据合并_*.csv',
            '/Users/admin/PycharmProjects/Weekly_ETF_Data_Changes/data/ETF_基础数据合并_*.csv',
            './data/ETF_基础数据合并_*.csv'
        ]
        
        for path in alt_paths:
            alt_files = glob.glob(path)
            if alt_files:
                files = alt_files
                print(f"在备选路径找到数据文件: {files}")
                break
                
    if not files:
        print("警告: 未找到ETF基础数据文件!")
        return "未找到ETF基础数据文件，请确保文件存在于正确位置"
    
    # 提取最新文件日期
    filename = sorted(files)[-1]
    print(f"找到最新数据文件: {filename}")
    date_str = filename.split('_')[-1].split('.')[0]
    
    try:
        # 读取ETF基础数据
        print(f"正在读取CSV文件: {filename}")
        etf_data = pd.read_csv(filename, encoding='utf-8-sig')
        print(f"成功读取数据，共 {len(etf_data)} 行")
        
        # 确保证券代码为字符串
        etf_data['证券代码'] = etf_data['证券代码'].astype(str).str.strip()
        
        # 获取当前日期
        current_date = datetime.strptime(date_str, "%Y%m%d")
        current_date_str = current_date.strftime("%m月%d日")
        
        # 计算上周日期
        previous_date = current_date - timedelta(days=7)
        previous_date_str = previous_date.strftime("%m月%d日")
        
        # 设置日期范围
        date_range = f"{previous_date.strftime('%m%d')}-{current_date.strftime('%m%d')}"
        
        # 读取商务协议文件
        商务协议_paths = [
            f'/Users/admin/Downloads/ETF单产品商务协议{date_str}.xlsx',
            f'./ETF单产品商务协议{date_str}.xlsx',
            f'./data/ETF单产品商务协议{date_str}.xlsx',
            f'/Users/admin/PycharmProjects/Weekly_ETF_Data_Changes/ETF单产品商务协议{date_str}.xlsx'
        ]
        
        商务协议_loaded = False
        for 商务协议_path in 商务协议_paths:
            try:
                print(f"尝试读取商务协议文件: {商务协议_path}")
                # 尝试读取商务协议文件
                商务协议 = pd.read_excel(商务协议_path, engine='openpyxl')
                print(f"成功读取商务协议文件，共 {len(商务协议)} 行")
                
                # 确保证券代码为字符串并去除空格
                商务协议['证券代码'] = 商务协议['证券代码'].astype(str).str.strip()
                
                # 添加商务品标记
                商务协议['是否商务品'] = '商务'
                
                # 合并商务协议数据
                etf_data = etf_data.merge(
                    商务协议[['证券代码', '是否商务品']],
                    on='证券代码',
                    how='left'
                ).fillna({'是否商务品': '非商务'})
                
                # 更新商务品集合
                business_etfs = set(商务协议['证券代码'].unique())
                商务协议_loaded = True
                print(f"商务品数据合并成功，共 {len(business_etfs)} 个商务品")
                break
                
            except FileNotFoundError:
                print(f"商务协议文件不存在: {商务协议_path}")
                continue
            except Exception as e:
                print(f"读取商务协议文件出错: {str(e)}")
                continue
        
        if not 商务协议_loaded:
            print("警告: 未能加载任何商务协议文件，将所有ETF标记为非商务品")
            etf_data['是否商务品'] = '非商务'
            business_etfs = set()
        
        # 确保关键列存在
        required_columns = [
            '关注人数变动', '持仓客户数变动', '保有金额变动',
            f'关注人数（{current_date_str}）', f'持仓客户数（{current_date_str}）'
        ]
        
        for col in required_columns:
            if col not in etf_data.columns:
                print(f"警告: 数据中缺少列 '{col}'，将创建默认值")
                etf_data[col] = 0
        
        print("数据加载和处理完成")
        return f"数据已加载，日期：{current_date_str}，商务品数量：{len(business_etfs)}个"
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"加载数据出错：{str(e)}")
        return f"加载数据出错：{str(e)}"

@app.route('/')
def index():
    """首页"""
    # 检查是否有code参数，如果有则预填充搜索框
    code = request.args.get('code', '')
    return render_template('dashboard.html', search_code=code)

# 搜索路由
@app.route('/search', methods=['POST'])
def search():
    """搜索ETF"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    code = request.form.get('code', '')
    
    if not code:
        return jsonify({"error": "请输入搜索关键词"})
    
    # 尝试不同的搜索方式
    # 1. 首先尝试按证券代码精确匹配
    matching_etfs = etf_data[etf_data['证券代码'] == code]
    
    if not matching_etfs.empty:
        # 找到了匹配的ETF证券代码
        tracking_index_code = matching_etfs.iloc[0]['跟踪指数代码'] if '跟踪指数代码' in matching_etfs.columns else None
        tracking_index_name = matching_etfs.iloc[0]['跟踪指数名称'] if '跟踪指数名称' in matching_etfs.columns else None
        
        # 查找同一指数的所有ETF
        if tracking_index_code:
            result_etfs = etf_data[etf_data['跟踪指数代码'] == tracking_index_code]
        else:
            result_etfs = matching_etfs
        
        search_type = "证券代码"
    else:
        # 2. 尝试按跟踪指数代码精确匹配
        matching_etfs = etf_data[etf_data['跟踪指数代码'] == code]
        
        if not matching_etfs.empty:
            # 找到了匹配的跟踪指数代码
            tracking_index_code = code
            tracking_index_name = matching_etfs.iloc[0]['跟踪指数名称'] if '跟踪指数名称' in matching_etfs.columns else None
            result_etfs = matching_etfs
            search_type = "跟踪指数代码"
        else:
            # 3. 尝试按跟踪指数名称进行模糊匹配
            if '跟踪指数名称' in etf_data.columns:
                matching_indices = etf_data[etf_data['跟踪指数名称'].str.contains(code, na=False, case=False)]
                
                if not matching_indices.empty:
                    # 找到了匹配的跟踪指数名称
                    # 按跟踪指数分组，并按规模排序
                    scale_col = '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'
                    if scale_col not in etf_data.columns:
                        # 尝试查找替代列
                        for col in etf_data.columns:
                            if '基金规模' in col and '亿元' in col:
                                scale_col = col
                                break
                    
                    # 按跟踪指数分组并计算总规模
                    if scale_col in etf_data.columns:
                        index_groups = []
                        # 获取匹配的跟踪指数列表
                        unique_indices = matching_indices['跟踪指数代码'].unique()
                        
                        for index_code in unique_indices:
                            if pd.isna(index_code):
                                continue
                                
                            index_etfs = etf_data[etf_data['跟踪指数代码'] == index_code]
                            index_name = index_etfs.iloc[0]['跟踪指数名称'] if '跟踪指数名称' in index_etfs.columns else "未知"
                            total_scale = index_etfs[scale_col].sum() if scale_col in index_etfs else 0
                            
                            index_groups.append({
                                'index_code': index_code,
                                'index_name': index_name,
                                'total_scale': total_scale,
                                'etfs': index_etfs
                            })
                        
                        # 按总规模排序
                        index_groups.sort(key=lambda x: x['total_scale'], reverse=True)
                        
                        # 构建结果
                        result = {
                            "success": True,
                            "search_type": "跟踪指数名称",
                            "keyword": code,
                            "index_groups": []
                        }
                        
                        for group in index_groups:
                            # 转换ETF数据为字典列表
                            etf_list = []
                            for _, row in group['etfs'].iterrows():
                                etf_dict = {}
                                for col in row.index:
                                    # 跳过NaN值
                                    if pd.notna(row[col]):
                                        etf_dict[col] = row[col]
                                etf_list.append(etf_dict)
                            
                            result["index_groups"].append({
                                "tracking_index_code": group['index_code'],
                                "tracking_index_name": group['index_name'],
                                "total_scale": round(group['total_scale'], 2),
                                "data": etf_list
                            })
                        
                        return jsonify(result)
                    else:
                        # 如果没有规模列，简单地按跟踪指数分组
                        result_etfs = matching_indices
                        tracking_index_code = None
                        tracking_index_name = code
                        search_type = "跟踪指数名称(无规模排序)"
                else:
                    # 4. 尝试按ETF产品名称进行模糊匹配
                    matching_etfs = etf_data[etf_data['产品名称'].str.contains(code, na=False, case=False)]
                    
                    if not matching_etfs.empty:
                        # 找到了匹配的ETF产品名称
                        # 获取第一个匹配的ETF的跟踪指数
                        tracking_index_code = matching_etfs.iloc[0]['跟踪指数代码'] if '跟踪指数代码' in matching_etfs.columns else None
                        tracking_index_name = matching_etfs.iloc[0]['跟踪指数名称'] if '跟踪指数名称' in matching_etfs.columns else None
                        result_etfs = matching_etfs
                        search_type = "ETF产品名称"
                    else:
                        return jsonify({"error": f"未找到与 '{code}' 相关的ETF或指数"})
            else:
                return jsonify({"error": "数据中缺少跟踪指数名称列，无法进行名称搜索"})
    
    # 如果不是按跟踪指数名称分组的情况，使用标准返回格式
    # 转换为字典列表
    result_list = []
    for _, row in result_etfs.iterrows():
        etf_dict = {}
        for col in row.index:
            # 跳过NaN值
            if pd.notna(row[col]):
                etf_dict[col] = row[col]
                
                # 对数值类型进行格式化
                if isinstance(etf_dict[col], (int, float)):
                    if '规模' in col or '亿' in col:
                        etf_dict[col] = round(etf_dict[col], 2)
                    elif '费率' in col:
                        etf_dict[col] = round(etf_dict[col], 4)
        result_list.append(etf_dict)
    
    return jsonify({
        "success": True,
        "search_type": search_type,
        "tracking_index_code": tracking_index_code,
        "tracking_index_name": tracking_index_name,
        "data": result_list
    })

# 添加缺失的路由

@app.route('/load_data')
def load_data_route():
    """加载数据的路由"""
    try:
        result = load_latest_data()
        if isinstance(result, str):
            # 如果返回的是错误消息字符串
            return jsonify({"success": False, "message": result})
        else:
            # 如果加载成功
            return jsonify({
                "success": True, 
                "message": f"数据已加载，日期：{current_date_str}，商务品数量：{len(business_etfs)}个"
            })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"加载数据出错：{str(e)}"})

@app.route('/overview')
def overview():
    """ETF市场概览"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 生成饼图 - 按基金管理人分布
        plt.figure(figsize=(10, 6))


        manager_counts = etf_data['基金管理人'].value_counts().head(10)
        
        # 使用简单的标签，避免中文问题
        labels = [f'公司{i+1}' for i in range(len(manager_counts))]
        wedges, texts, autotexts = plt.pie(manager_counts, labels=labels, autopct='%1.1f%%', startangle=90)
        
        # 添加图例，显示实际的公司名称
        plt.legend(wedges, manager_counts.index, title="基金管理人", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        
        plt.title('ETF基金管理人分布（Top 10）')
        plt.axis('equal')
        
        # 将图表转换为base64编码
        pie_buffer = io.BytesIO()
        plt.savefig(pie_buffer, format='png', dpi=100, bbox_inches='tight')
        pie_buffer.seek(0)
        pie_chart = base64.b64encode(pie_buffer.getvalue()).decode('utf-8')
        plt.close()
        
        # 生成柱状图 - 按基金规模
        plt.figure(figsize=(12, 6))
        
        # 确保列名存在
        scale_col = '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'
        if scale_col not in etf_data.columns:
            # 尝试查找替代列
            for col in etf_data.columns:
                if '基金规模' in col and '亿元' in col:
                    scale_col = col
                    break
        
        # 按管理人分组并计算规模总和
        if scale_col in etf_data.columns:
            company_scale = etf_data.groupby('基金管理人')[scale_col].sum().sort_values(ascending=False).head(10)
            
            # 使用条形图而不是company_scale.plot
            plt.bar(range(len(company_scale)), company_scale.values)
            plt.xticks(range(len(company_scale)), company_scale.index, rotation=45)
            plt.title('ETF基金管理人规模排名（Top 10）')
            plt.xlabel('基金管理人')
            plt.ylabel('规模（亿元）')
            plt.tight_layout()
            
            # 将图表转换为base64编码
            company_buffer = io.BytesIO()
            plt.savefig(company_buffer, format='png', dpi=100, bbox_inches='tight')
            company_buffer.seek(0)
            company_chart = base64.b64encode(company_buffer.getvalue()).decode('utf-8')
            plt.close()
        else:
            company_chart = None
        
        # 统计数据
        total_etfs = len(etf_data)
        total_companies = etf_data['基金管理人'].nunique()
        
        # 计算总规模
        if scale_col in etf_data.columns:
            total_scale = etf_data[scale_col].sum()
        else:
            total_scale = None
        
        # 返回结果
        return jsonify({
            "success": True,
            "stats": {
                "total_etfs": total_etfs,
                "total_companies": total_companies,
                "total_scale": round(total_scale, 2) if total_scale is not None else "未知",
                "business_etfs": len(business_etfs)
            },
            "charts": {
                "pie_chart": pie_chart,
                "company_chart": company_chart
            }
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"生成市场概览出错：{str(e)}"})

@app.route('/business_analysis')
def business_analysis():
    """商务品分析"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 筛选商务品ETF
        business_df = etf_data[etf_data['证券代码'].isin(business_etfs)]
        
        if business_df.empty:
            return jsonify({"error": "未找到商务品ETF数据"})
        
        # 按管理人分组统计商务品数量
        business_by_company = business_df.groupby('基金管理人').size().sort_values(ascending=False)
        
        # 统计数据
        total_business = len(business_df)
        total_companies = business_df['基金管理人'].nunique()
        
        # 确保列名存在
        scale_col = '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'
        if scale_col not in business_df.columns:
            # 尝试查找替代列
            for col in business_df.columns:
                if '基金规模' in col and '亿元' in col:
                    scale_col = col
                    break
        
        # 计算总规模
        if scale_col in business_df.columns:
            total_scale = business_df[scale_col].sum()
            # 按管理人分组计算规模
            scale_by_company = business_df.groupby('基金管理人')[scale_col].sum().sort_values(ascending=False)
        else:
            total_scale = None
            scale_by_company = None
        
        # 构建结果
        result = {
            "success": True,
            "stats": {
                "total_business": total_business,
                "total_companies": total_companies,
                "total_scale": round(total_scale, 2) if total_scale is not None else "未知"
            },
            "by_company": []
        }
        
        # 添加公司数据
        for company in business_by_company.index:
            company_data = {
                "company": company,
                "count": int(business_by_company[company])
            }
            
            if scale_by_company is not None and company in scale_by_company:
                company_data["scale"] = round(float(scale_by_company[company]), 2)
            
            result["by_company"].append(company_data)
        
        return jsonify(result)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"商务品分析出错：{str(e)}"})

@app.route('/generate_report')
def generate_report():
    """生成报告"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 创建一个简单的Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f'ETF市场周报 ({current_date_str})', 0)
        
        # 添加市场概览
        doc.add_heading('市场概览', level=1)
        doc.add_paragraph(f'截至{current_date_str}，市场上共有{len(etf_data)}只ETF产品，其中商务品{len(business_etfs)}只。')
        
        # 添加商务品分析
        doc.add_heading('商务品分析', level=1)
        business_df = etf_data[etf_data['证券代码'].isin(business_etfs)]
        
        # 按管理人分组统计商务品数量
        business_by_company = business_df.groupby('基金管理人').size().sort_values(ascending=False)
        
        # 添加表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '基金管理人'
        header_cells[1].text = '商务品数量'
        
        # 添加数据行
        for company, count in business_by_company.items():
            row_cells = table.add_row().cells
            row_cells[0].text = company
            row_cells[1].text = str(count)
        
        # 保存文档
        report_filename = f'ETF市场周报_{current_date_str}.docx'
        report_path = os.path.join(app.config['UPLOAD_FOLDER'], report_filename)
        doc.save(report_path)
        
        return jsonify({
            "success": True,
            "message": "报告生成成功",
            "report_url": f"/download_report/{report_filename}"
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"生成报告出错：{str(e)}"})

@app.route('/download_report/<filename>')
def download_report(filename):
    """下载报告"""
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename), as_attachment=True)

@app.route('/upload_data', methods=['GET', 'POST'])
def upload_data():
    """上传数据文件"""
    if request.method == 'POST':
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({"error": "未选择文件"})
        
        file = request.files['file']
        
        # 检查文件名
        if file.filename == '':
            return jsonify({"error": "未选择文件"})
        
        if file:
            # 安全地获取文件名
            filename = secure_filename(file.filename)
            
            # 保存文件
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            return jsonify({
                "success": True,
                "message": f"文件 {filename} 上传成功",
                "file_path": file_path
            })
    
    return render_template('upload.html')

# 创建templates目录
os.makedirs('templates', exist_ok=True)

# 创建upload.html模板
with open('templates/upload.html', 'w', encoding='utf-8') as f:
    f.write('''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>上传ETF数据文件</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {
            padding: 20px;
        }
        .upload-container {
            max-width: 600px;
            margin: 0 auto;
        }
        .card {
            margin-top: 20px;
            box-shadow: 0 0.125rem 0.25rem rgba(0, 0, 0, 0.075);
        }
    </style>
</head>
<body>
    <div class="container upload-container">
        <h2 class="text-center mb-4">上传ETF数据文件</h2>
        
        <div class="card">
            <div class="card-body">
                <form action="/upload_data" method="post" enctype="multipart/form-data">
                    <div class="mb-3">
                        <label for="file" class="form-label">选择Excel文件</label>
                        <input type="file" class="form-control" id="file" name="file" accept=".xlsx,.xls,.csv">
                        <div class="form-text">支持Excel格式(.xlsx, .xls)和CSV格式(.csv)的文件</div>
                    </div>
                    <div class="d-grid gap-2">
                        <button type="submit" class="btn btn-primary">上传文件</button>
                        <a href="/" class="btn btn-secondary">返回首页</a>
                    </div>
                </form>
            </div>
        </div>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
''')

def load_etf_data():
    """加载ETF数据"""
    global etf_data, business_etfs, current_date_str, previous_date_str, date_range
    
    # 查找最新的Excel文件
    data_files = glob.glob(os.path.join(app.config['UPLOAD_FOLDER'], '*.xlsx'))
    
    if not data_files:
        return False, "未找到ETF数据文件，请先上传数据"
    
    # 按修改时间排序，获取最新的文件
    latest_file = max(data_files, key=os.path.getmtime)
    
    try:
        # 读取Excel文件
        etf_data = pd.read_excel(latest_file)
        
        # 提取日期信息
        file_name = os.path.basename(latest_file)
        date_parts = file_name.split('_')
        if len(date_parts) >= 2:
            try:
                current_date = datetime.strptime(date_parts[0], '%Y%m%d')
                current_date_str = current_date.strftime('%Y年%m月%d日')
                
                # 计算上一周日期
                previous_date = current_date - timedelta(days=7)
                previous_date_str = previous_date.strftime('%Y年%m月%d日')
                
                # 日期范围
                date_range = f"{previous_date.strftime('%m%d')}-{current_date.strftime('%m%d')}"
            except:
                current_date_str = "未知日期"
                previous_date_str = "未知日期"
                date_range = "未知日期范围"
        
        # 标记商务品
        if '是否商务品' not in etf_data.columns:
            etf_data['是否商务品'] = '非商务'
        
        # 提取商务品代码集合
        business_etfs = set(etf_data[etf_data['是否商务品'] == '商务']['证券代码'].tolist())
        
        # 处理管理人简称
        if '基金管理人' in etf_data.columns:
            etf_data['管理人'] = etf_data['基金管理人'].apply(get_manager_short)
        
        # 确保必要的列存在
        required_columns = ['证券代码', '产品名称', '基金管理人', '管理人', '跟踪指数代码', '跟踪指数名称', 
                           '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元', 
                           '管理费率(%)', '月日均交易额(百万)', '是否商务品']
        
        for col in required_columns:
            if col not in etf_data.columns and col not in ['管理人']:
                return False, f"数据文件缺少必要的列：{col}"
        
        return True, f"成功加载ETF数据，日期范围：{date_range}"
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, f"加载数据出错：{str(e)}"

if __name__ == '__main__':
    # 确保目录存在
    os.makedirs('templates', exist_ok=True)
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], 'reports'), exist_ok=True)
    
    # 创建data目录用于存放数据文件
    os.makedirs('data', exist_ok=True)
    
    # 启动应用前自动加载数据
    print("正在预加载ETF数据...")
    load_result = load_latest_data()
    print(f"数据加载结果: {load_result}")
    
    # 启动应用
    app.run(debug=True, host='0.0.0.0', port=5001)