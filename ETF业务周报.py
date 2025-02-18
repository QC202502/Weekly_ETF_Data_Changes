import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re  # 新增正则表达式模块导入

def preprocess_data():
    # 数据预处理
    result = pd.read_csv('ETF_基础数据合并_20250214.csv', encoding='utf-8-sig')
    商务协议 = pd.read_excel('/Users/admin/Downloads/ETF单产品商务协议20250214.xlsx', engine='openpyxl')

    # 列名日期提取（保持原有数据处理流程）
    date_pattern = re.compile(r'.*（(\d+)月(\d+)日）')
    dates = set()
    for col in result.columns:
        match = date_pattern.search(col)
        if match:
            month = f"{int(match.group(1)):02d}"
            day = f"{int(match.group(2)):02d}"
            dates.add(f"{month}{day}")

    # 合并商务协议数据（原有逻辑）
    result['证券代码'] = result['证券代码'].astype(str).str.strip()
    商务协议['证券代码'] = 商务协议['证券代码'].astype(str).str.strip()
    商务协议['是否商务品'] = '商务'
    result = result.merge(
        商务协议[['证券代码', '是否商务品']],
        on='证券代码',
        how='left'
    ).fillna({'是否商务品': '非商务'})

    # 填充空值（原有逻辑）
    result['基金管理人'] = result['基金管理人'].fillna('未知管理人')
    result['跟踪指数名称'] = result['跟踪指数名称'].fillna('未知指数')

    # 数值处理（原有逻辑）
    numeric_cols = ['关注人数（02月14日）', '关注人数（02月07日）',
                    '持仓客户数（02月14日）', '持仓客户数（02月07日）',
                    '保有金额（02月14日）', '保有金额（02月07日）']
    result[numeric_cols] = result[numeric_cols].fillna(0).astype(int)

    # 计算变动（原有逻辑）
    for metric in ['关注人数', '持仓客户数', '保有金额']:
        result[f'{metric}变动'] = result[f'{metric}（02月14日）'] - result[f'{metric}（02月07日）']

    # 最终返回两个值（必须在所有数据处理完成后返回）
    return result, '-'.join(sorted(dates))  # 返回处理后的数据和日期范围


def get_manager_short(full_name):
    """基金管理人简称提取"""
    special_mapping = {
        '汇添富基金管理有限公司': '汇添富',
        '易方达基金管理有限公司': '易方达',
        '华夏基金管理有限公司': '华夏',
        '南方基金管理股份有限公司': '南方',
        '嘉实基金管理有限公司': '嘉实'
    }
    if full_name in special_mapping:
        return special_mapping[full_name]

    patterns = ['基金管理有限公司', '基金管理公司', '基金', '管理有限公司', '股份有限公司']
    for pattern in patterns:
        if pattern in full_name:
            return full_name.split(pattern)[0]

    return full_name[:4] if len(full_name) >= 4 else full_name


class ETFReporter:
    def __init__(self, data, date_range):
        self.doc = Document()
        self.data = data
        self.date_range = date_range  # 新增实例变量存储日期范围

        # 解析动态日期范围
        start_str, end_str = date_range.split('-')

        # 解析起始日期
        start_month = int(start_str[:2])
        start_day = int(start_str[2:])
        self.start_date = f"{start_month}月{start_day}日"

        # 解析结束日期（解决file_date缺失问题）
        end_month = int(end_str[:2])
        end_day = int(end_str[2:])
        self.end_date = f"{end_month}月{end_day}日"
        self.file_date = self.end_date  # 保持向后兼容

        self._set_styles()

    def _set_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        font.size = Pt(10.5)

    def generate_report(self):
        self._add_cover()
        self._add_combined_analysis('关注人数', '关注人次', {'total': '', 'delta': ''})
        self._add_combined_analysis('持仓客户数', '持仓客户数', {'total': '户', 'delta': '户'})
        self._add_combined_analysis('保有金额', '保有金额', {'total': '亿', 'delta': '百万'})  # 修改delta单位
        self._add_tables()
        self.doc.save(f'ETF产品运营周报（{self.date_range}）.docx')  # 修复文件名生成
        print("周报生成完成！")

    def _add_cover(self):
        self.doc.add_heading('ETF产品运营周报', 0)
        p = self.doc.add_paragraph()
        # 动态日期范围（修改此处）
        date_range_str = f'报告时间范围：2025年{self.start_date}至{self.end_date}\n'
        p.add_run(date_range_str).bold = True
        p.add_run('汇报人：邱超').italic = True
    def _add_combined_analysis(self, metric, name, unit_map):
        total_data = self.data
        total = total_data[f'{metric}（02月14日）'].sum()
        delta = total_data[f'{metric}变动'].sum()

        biz_data = self.data[self.data['是否商务品'] == '商务']
        biz_total = biz_data[f'{metric}（02月14日）'].sum()
        biz_delta = biz_data[f'{metric}变动'].sum()

        self.doc.add_heading(f"{name}分析", level=2)
        p = self.doc.add_paragraph()

        # 总体描述（加粗）
        total_str = self._format_value(total, metric, is_total=True) + unit_map['total']
        delta_str = self._format_delta(delta, unit_map['delta'], metric)
        main_text = f"截止{self.end_date}，{name}总量为{total_str}，本周变动：{delta_str}。"
        run = p.add_run(main_text)
        run.bold = True

        self._add_top_bottom(p, total_data, metric)

        if not biz_data.empty:
            # 商务部分（加粗）
            biz_total_str = self._format_value(biz_total, metric, is_total=True) + unit_map['total']
            biz_delta_str = self._format_delta(biz_delta, unit_map['delta'], metric)

            biz_p = self.doc.add_paragraph()
            biz_text = f"其中，商务品{name}总量为{biz_total_str}，本周变动：{biz_delta_str}。"
            biz_run = biz_p.add_run(biz_text)
            biz_run.bold = True

            self._add_top_bottom(biz_p, biz_data, metric, is_business=True)


    # 修改_format_value方法（解决第二个问题）
    def _format_value(self, value, metric, is_total=False):
        """统一数值格式化"""
        if metric == '保有金额':
            if is_total:  # 总量用亿
                return f"{value / 1e8:,.2f}"
            else:  # 变动值根据量级自动转换
                if abs(value) >= 1e8:
                    return f"{value / 1e8:,.2f}亿"
                elif abs(value) >= 1e4:
                    return f"{value / 1e4:,.0f}万"
                else:
                    return f"{value:,.0f}元"
        return f"{value:,}"

    # 修改_format_delta方法（解决第一个问题）
    def _format_delta(self, delta, unit, metric):
        """智能单位转换"""
        sign = '+' if delta >= 0 else '-'
        abs_value = abs(delta)

        # 特殊处理保有金额
        if metric == '保有金额':
            if abs_value >= 1e8:  # 超过1亿
                value = abs_value / 1e8
                return f"{sign}{value:,.2f}亿"
            elif abs_value >= 1e4:  # 超过1万
                value = abs_value / 1e4
                return f"{sign}{value:,.0f}万"
            else:
                return f"{sign}{abs_value:,.0f}元"
        else:
            return f"{sign}{self._format_value(abs_value, metric)}{unit}"
    def _add_top_bottom(self, paragraph, data, metric, is_business=False):
        top3 = data.nlargest(3, f'{metric}变动')
        bottom3 = data.nsmallest(3, f'{metric}变动')

        increase_str = self._format_products(top3, metric, is_business)
        decrease_str = self._format_products(bottom3, metric, is_business)

        paragraph.add_run(f"增加前三：{increase_str}。")
        paragraph.add_run(f"减少前三：{decrease_str}。")

    # 修改_format_products方法（解决第一个问题）
    def _format_products(self, df, metric, is_business):
        """产品条目格式化（带符号）"""
        items = []
        for _, row in df.iterrows():
            manager = get_manager_short(row['基金管理人'])
            delta = row[f'{metric}变动']

            # 获取带符号的变动值
            sign = '+' if delta >= 0 else '-'
            abs_delta = abs(delta)

            # 特殊处理保有金额
            if metric == '保有金额':
                if abs_delta >= 1e8:
                    delta_str = f"{sign}{abs_delta / 1e8:,.2f}亿"
                elif abs_delta >= 1e4:
                    delta_str = f"{sign}{abs_delta / 1e4:,.0f}万"
                else:
                    delta_str = f"{sign}{abs_delta:,.0f}元"
            else:
                delta_str = f"{sign}{self._format_value(abs_delta, metric)}"

            # 构造条目
            components = [delta_str, manager]
            if not is_business:
                components.append(row['是否商务品'])

            item_str = f"{row['证券名称（02月14日）']} ({'、'.join(components)})"
            items.append(item_str)
        return "，".join(items)

    def _add_tables(self):
        for metric in ['关注人数', '持仓客户数', '保有金额']:
            top20 = self.data.nlargest(20, f'{metric}变动')
            self._add_table(f'{metric}增长Top20', top20, metric)

            bottom20 = self.data.nsmallest(20, f'{metric}变动')
            self._add_table(f'{metric}减少Top20', bottom20, metric)

    def _add_table(self, title, data, metric):
        self.doc.add_heading(title, level=2)
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading Accent 1'

        headers = ['证券代码', '产品名称', '管理人', '当前值', '上周值', '变动值', '商务属性']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        for _, row in data.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['证券代码'])
            cells[1].text = row['证券名称（02月14日）']
            cells[2].text = get_manager_short(row['基金管理人'])
            # 修复4：使用类方法并传递参数
            cells[3].text = self._format_value(row[f'{metric}（02月14日）'], metric, is_total=True)
            cells[4].text = self._format_value(row[f'{metric}（02月07日）'], metric, is_total=True)
            delta_value = row[f'{metric}变动']
            cells[5].text = self._format_delta(delta_value, '', metric)
            cells[6].text = row['是否商务品']


if __name__ == "__main__":
    # 获取预处理数据和日期范围
    df, date_range = preprocess_data()  # 解包两个返回值

    # 传入date_range参数
    reporter = ETFReporter(df, date_range)
    reporter.generate_report()