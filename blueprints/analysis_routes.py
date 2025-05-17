from flask import Blueprint, jsonify
import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import io
import base64
import pandas as pd
from docx import Document
import os
import logging # 添加导入
import numpy as np # 需要导入 numpy 来检查 np.isnan

# 配置 logger
logger = logging.getLogger(__name__) # 添加 logger 初始化

# 创建蓝图
analysis_bp = Blueprint('analysis', __name__)

# 从数据蓝图导入全局变量
from blueprints.data_routes import etf_data, business_etfs, current_date_str

@analysis_bp.route('/overview')
def overview():
    """ETF市场概览"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 生成饼图 - 按基金管理人分布
        plt.figure(figsize=(10, 6))

        manager_counts = etf_data['基金管理人'].value_counts().head(10)
        
        # 使用简单的标签，避免中文问题
        labels = [f'公司{i+1}' for i in range(len(manager_counts))]
        wedges, texts, autotexts = plt.pie(manager_counts, labels=labels, autopct='%1.1f%%', startangle=90)
        
        # 添加图例，显示实际的公司名称
        plt.legend(wedges, manager_counts.index, title="基金管理人", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        
        plt.title('ETF基金管理人分布（Top 10）')
        plt.axis('equal')
        
        # 将图表转换为base64编码
        pie_buffer = io.BytesIO()
        plt.savefig(pie_buffer, format='png', dpi=100, bbox_inches='tight')
        pie_buffer.seek(0)
        pie_chart = base64.b64encode(pie_buffer.getvalue()).decode('utf-8')
        plt.close()
        
        # 生成柱状图 - 按基金规模
        plt.figure(figsize=(12, 6))
        
        # 确保列名存在
        scale_col = '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'
        if scale_col not in etf_data.columns:
            # 尝试查找替代列
            for col in etf_data.columns:
                if '基金规模' in col and '亿元' in col:
                    scale_col = col
                    break
        
        # 按管理人分组并计算规模总和
        if scale_col in etf_data.columns:
            company_scale = etf_data.groupby('基金管理人')[scale_col].sum().sort_values(ascending=False).head(10)
            
            # 使用条形图而不是company_scale.plot
            plt.bar(range(len(company_scale)), company_scale.values)
            plt.xticks(range(len(company_scale)), company_scale.index, rotation=45)
            plt.title('ETF基金管理人规模排名（Top 10）')
            plt.xlabel('基金管理人')
            plt.ylabel('规模（亿元）')
            plt.tight_layout()
            
            # 将图表转换为base64编码
            company_buffer = io.BytesIO()
            plt.savefig(company_buffer, format='png', dpi=100, bbox_inches='tight')
            company_buffer.seek(0)
            company_chart = base64.b64encode(company_buffer.getvalue()).decode('utf-8')
            plt.close()
        else:
            company_chart = None
        
        # 统计数据
        total_etfs = len(etf_data)
        total_companies = etf_data['基金管理人'].nunique()
        
        # 计算总规模
        if scale_col in etf_data.columns:
            total_scale = etf_data[scale_col].sum()
        else:
            total_scale = None
        
        # 返回结果
        return jsonify({
            "success": True,
            "stats": {
                "total_etfs": total_etfs,
                "total_companies": total_companies,
                "total_scale": round(total_scale, 2) if total_scale is not None else "未知",
                "business_etfs": len(business_etfs)
            },
            "charts": {
                "pie_chart": pie_chart,
                "company_chart": company_chart
            }
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"生成市场概览出错：{str(e)}"})

@analysis_bp.route('/business_analysis')
def business_analysis():
    """商务品分析"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 筛选商务品ETF
        business_df = etf_data[etf_data['证券代码'].isin(business_etfs)]
        
        if business_df.empty:
            return jsonify({"error": "未找到商务品ETF数据"})
        
        # 按管理人分组统计商务品数量
        business_by_company = business_df.groupby('基金管理人').size().sort_values(ascending=False)
        
        # 统计数据
        total_business = len(business_df)
        total_companies = business_df['基金管理人'].nunique()
        
        # 确保列名存在
        scale_col = '基金规模(合计)[交易日期]S_cal_date(now(),0,D,0)[单位]亿元'
        if scale_col not in business_df.columns:
            # 尝试查找替代列
            for col in business_df.columns:
                if '基金规模' in col and '亿元' in col:
                    scale_col = col
                    break
        
        # 计算总规模
        if scale_col in business_df.columns:
            total_scale = business_df[scale_col].sum()
            # 按管理人分组计算规模
            scale_by_company = business_df.groupby('基金管理人')[scale_col].sum().sort_values(ascending=False)
        else:
            total_scale = None
            scale_by_company = None
        
        # 构建结果
        result = {
            "success": True,
            "stats": {
                "total_business": total_business,
                "total_companies": total_companies,
                "total_scale": round(total_scale, 2) if total_scale is not None else "未知"
            },
            "by_company": []
        }
        
        # 添加公司数据
        for company in business_by_company.index:
            company_data = {
                "company": company,
                "count": int(business_by_company[company])
            }
            
            if scale_by_company is not None and company in scale_by_company:
                company_data["scale"] = round(float(scale_by_company[company]), 2)
            
            result["by_company"].append(company_data)
        
        return jsonify(result)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"商务品分析出错：{str(e)}"})

@analysis_bp.route('/generate_report')
def generate_report():
    """生成报告"""
    if etf_data is None:
        return jsonify({"error": "数据未加载，请先加载数据"})
    
    try:
        # 创建一个简单的Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f'ETF市场周报 ({current_date_str})', 0)
        
        # 添加市场概览
        doc.add_heading('市场概览', level=1)
        doc.add_paragraph(f'截至{current_date_str}，市场上共有{len(etf_data)}只ETF产品，其中商务品{len(business_etfs)}只。')
        
        # 添加商务品分析
        doc.add_heading('商务品分析', level=1)
        business_df = etf_data[etf_data['证券代码'].isin(business_etfs)]
        
        # 按管理人分组统计商务品数量
        business_by_company = business_df.groupby('基金管理人').size().sort_values(ascending=False)
        
        # 添加表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '基金管理人'
        header_cells[1].text = '商务品数量'
        
        # 添加数据行
        for company, count in business_by_company.items():
            row_cells = table.add_row().cells
            row_cells[0].text = company
            row_cells[1].text = str(count)
        
        # 保存文档
        from flask import current_app
        report_filename = f'ETF市场周报_{current_date_str}.docx'
        report_path = os.path.join(current_app.config['UPLOAD_FOLDER'], report_filename)
        doc.save(report_path)
        
        return jsonify({
            "success": True,
            "message": "报告生成成功",
            "report_url": f"/download_report/{report_filename}"
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"生成报告出错：{str(e)}"})

@analysis_bp.route('/api/business_data')
def api_business_data():
    """提供商务品分析散点图所需的数据。

    返回每个公司的:
    - company_short_name (公司简称)
    - total_holding_value (总持仓规模, 作为 x)
    - total_business_value (商务品总规模, 作为 y)
    - business_agreement_ratio (商务品占比, 作为 ratio)
    - product_count (产品数量, 用于点大小)
    """
    try:
        from database.models import Database # 确保导入Database类
        db = Database()
        # 假设 get_company_analytics_for_dashboard 可以返回我们需要的所有字段
        # 或者我们需要一个新的方法。这里我们先尝试使用现有的，如果字段不足后续再调整。
        # 我们需要的数据字段：company_short_name, total_holding_value, 
        # total_business_value, business_agreement_ratio, product_count
        
        # 尝试直接从 etf_company_analytics 获取原始数据，因为JS端会做筛选和TOP N处理
        # 如果直接调用 get_company_analytics_for_dashboard，它内部可能有排序和限制
        conn = db.connect()
        query = """
            SELECT 
                company_short_name, 
                total_holding_value, 
                business_total_holding_value,
                business_agreement_ratio, 
                product_count 
            FROM etf_company_analytics
        """
        # 使用 pandas 读取数据，方便转换为字典列表
        df = pd.read_sql_query(query, conn)
        conn.close()

        if df.empty:
            return jsonify({"error": "未能从数据库获取商务品分析数据"}), 500

        # 将列名转换为JS期望的格式
        df.rename(columns={
            'company_short_name': 'company',
            'total_holding_value': 'x',
            'business_total_holding_value': 'y',
            'business_agreement_ratio': 'ratio',
            'product_count': 'productCount'
        }, inplace=True)

        # 确保所有数值列都是float或int，并且处理NaN/None为0，避免JSON序列化问题
        for col in ['x', 'y', 'ratio', 'productCount']:
            if col in df.columns:
                # 确保先转换为数值，再fillna
                df[col] = pd.to_numeric(df[col], errors='coerce')
                df[col] = df[col].fillna(0)
            else:
                # 如果某个期望的列不存在，用0填充，但前端JS需要能处理
                df[col] = 0 
                logger.warning(f"列 {col} 在商务品分析数据中缺失，已用0填充。")
        
        # 转换ratio为百分比数值 (例如 0.75 -> 75)
        if 'ratio' in df.columns: # 检查列是否存在
            # 先确保是数值类型
            df['ratio'] = pd.to_numeric(df['ratio'], errors='coerce').fillna(0)
            # 检查最大值，避免已经是百分比的再次乘以100
            # 只有当ratio的最大值在 (0, 1] 区间时才乘以100
            if not df['ratio'].empty:
                max_ratio = df['ratio'].max()
                if 0 < max_ratio <= 1.0:
                    df['ratio'] = df['ratio'] * 100
                elif max_ratio > 100: # 如果数值已经很大，可能本身就是百分比，但单位不对，这里可能需要进一步的逻辑或警告
                    logger.warning(f"Ratio列的最大值 {max_ratio} 似乎已经大于100，请检查数据源是否已经是百分比但未正确处理。")
            # 如果列为空或所有值都为0，则max_ratio会是0，不会进入乘100的逻辑，这是正确的

        chart_data = df.to_dict(orient='records')
        
        # 还需要返回一些汇总统计数据，如 modules/business.html 中顶部的徽章所示
        # total_business: 拥有商务品的公司数量 (这里可以理解为 chart_data 中 y > 0 的公司数)
        # business_companies: 基金公司总数 (即 chart_data 的长度)
        # business_scale: 商务品总持仓规模 (即所有公司 y 值的总和)

        total_business_companies_with_products = df[df['y'] > 0].shape[0]
        total_companies_count = len(chart_data)
        
        # 明确处理 total_business_scale 可能为 NaN 的情况
        raw_total_business_scale = df['y'].sum()
        if pd.isna(raw_total_business_scale): # 使用 pd.isna() 来检查NaN
            total_business_scale = 0.0
        else:
            # 确保转换为Python float类型，因为numpy float类型有时也不能直接被jsonify处理
            total_business_scale = float(raw_total_business_scale)

        return jsonify({
            "success": True,
            "chart_data": chart_data,
            "stats": {
                "total_business": total_business_companies_with_products, # 持有商务品的公司数
                "business_companies": total_companies_count, # 基金公司总数 (参与图表)
                "business_scale": round(total_business_scale, 2)
            }
        })

    except Exception as e:
        logger.error(f"获取商务品分析API数据出错: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({"error": f"服务器内部错误: {str(e)}"}), 500