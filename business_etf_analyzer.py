import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import glob
from datetime import datetime, timedelta
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import numpy as np

# 版本信息
__version__ = "1.0.0"   
RELEASE_DATE = "2025-03-15"

def show_version():
    """显示版本信息"""
    print(f"ETF 商务品分析工具 Version: {__version__} ({RELEASE_DATE})")
    print("Copyright © 2024 邱超. All rights reserved.")

def load_data():
    """加载并预处理数据"""
    # 动态获取最新数据文件
    files = glob.glob('ETF_基础数据合并_*.csv')
    if not files:
        raise FileNotFoundError("未找到ETF基础数据文件")
    
    # 提取最新文件日期
    filename = sorted(files)[-1]
    date_str = filename.split('_')[-1].split('.')[0]
    
    try:
        current_date = datetime.strptime(date_str, "%Y%m%d")
    except ValueError:
        raise ValueError("文件名日期格式应为ETF_基础数据合并_YYYYMMDD.csv")
    
    previous_date = current_date - timedelta(days=7)
    
    # 修正日期格式（保留前导零）
    current_str = current_date.strftime("%m月%d日")
    previous_str = previous_date.strftime("%m月%d日")

    # 读取数据
    result = pd.read_csv(filename, encoding='utf-8-sig')
    
    # 读取ETF指数分类数据
    classification_path = f'/Users/admin/Downloads/ETF-Index-Classification_{date_str}.xlsx'
    try:
        classification_data = pd.read_excel(classification_path, engine='openpyxl')
        print(f"指数分类文件路径: {classification_path}")
        print(f"指数分类文件总行数: {len(classification_data)}")
        
        # 合并分类数据
        if '跟踪指数代码' in result.columns and '跟踪指数代码' in classification_data.columns:
            result = pd.merge(
                result,
                classification_data,
                on='跟踪指数代码',
                how='left'
            )
            print(f"成功合并指数分类数据")
        else:
            print(f"警告: 无法合并指数分类数据，缺少'跟踪指数代码'列")
    except FileNotFoundError:
        print(f"警告: 指数分类文件不存在: {classification_path}，尝试使用固定路径")
        # 尝试使用固定路径
        classification_path = '/Users/admin/Downloads/ETF-Index-Classification_20240307.xlsx'
        try:
            classification_data = pd.read_excel(classification_path, engine='openpyxl')
            print(f"使用固定路径指数分类文件: {classification_path}")
            print(f"指数分类文件总行数: {len(classification_data)}")
            
            # 合并分类数据
            if '跟踪指数代码' in result.columns and '跟踪指数代码' in classification_data.columns:
                result = pd.merge(
                    result,
                    classification_data,
                    on='跟踪指数代码',
                    how='left'
                )
                print(f"成功合并指数分类数据")
            else:
                print(f"警告: 无法合并指数分类数据，缺少'跟踪指数代码'列")
        except Exception as e:
            print(f"警告: 读取指数分类文件失败: {str(e)}")
    except Exception as e:
        print(f"警告: 读取指数分类文件失败: {str(e)}")
    
    # 动态商务协议文件
    商务协议_path = f'/Users/admin/Downloads/ETF单产品商务协议{date_str}.xlsx'
    try:
        商务协议 = pd.read_excel(商务协议_path, engine='openpyxl')
        
        # 详细检查商务协议文件
        print(f"商务协议文件路径: {商务协议_path}")
        print(f"商务协议文件总行数: {len(商务协议)}")
        
        # 检查是否有重复的证券代码
        重复代码 = 商务协议['证券代码'].duplicated().sum()
        print(f"商务协议中重复的证券代码数量: {重复代码}")
        
        if 重复代码 > 0:
            重复代码列表 = 商务协议[商务协议['证券代码'].duplicated(keep=False)]['证券代码'].unique()
            print(f"重复的证券代码列表: {list(重复代码列表)}")
            
            # 显示每个重复代码的详细信息
            print("重复证券代码的详细信息:")
            for code in 重复代码列表:
                重复项 = 商务协议[商务协议['证券代码'] == code]
                print(f"证券代码 {code} 出现 {len(重复项)} 次:")
                for idx, row in 重复项.iterrows():
                    print(f"  - 行 {idx+1}: {row['产品名称']} ({row['基金公司简称']})")
        
        # 确保证券代码为字符串并去除空格
        商务协议['证券代码'] = 商务协议['证券代码'].astype(str).str.strip()
        
        # 再次检查唯一证券代码数量
        唯一代码数量 = len(商务协议['证券代码'].unique())
        print(f"商务协议中唯一证券代码数量: {唯一代码数量}")
        
    except FileNotFoundError:
        raise FileNotFoundError(f"商务协议文件缺失：{商务协议_path}")
    except Exception as e:
        raise Exception(f"读取商务协议文件出错: {str(e)}")

    # 检查商务协议文件中的列
    print(f"商务协议文件列名: {list(商务协议.columns)}")
    
    # 确定要合并的列
    merge_columns = ['证券代码', '是否商务品', '产品名称', '基金公司简称']
    
    # 检查可选列是否存在，如果存在则添加到合并列中
    optional_columns = ['合作开始日期', '合作结束日期', '分成比例']
    for col in optional_columns:
        if col in 商务协议.columns:
            merge_columns.append(col)
        else:
            print(f"警告: 商务协议文件中缺少'{col}'列，将使用默认值")
            # 为缺失的列添加默认值
            if col == '分成比例':
                商务协议[col] = 0.4  # 默认分成比例为40%
            else:
                商务协议[col] = None
    
    # 合并商务协议数据前的处理
    result['证券代码'] = result['证券代码'].astype(str).str.strip()
    商务协议['是否商务品'] = '商务'
    
    # 检查合并前后的商务品数量
    商务品代码集合 = set(商务协议['证券代码'].unique())
    基础数据代码集合 = set(result['证券代码'].unique())
    未匹配商务品 = 商务品代码集合 - 基础数据代码集合
    
    if len(未匹配商务品) > 0:
        print(f"警告: 有 {len(未匹配商务品)} 个商务品在基础数据中未找到匹配项")
        print(f"未匹配的商务品代码: {sorted(list(未匹配商务品))}")
    
    # 合并商务协议数据
    result = result.merge(
        商务协议[merge_columns],
        on='证券代码',
        how='left'
    ).fillna({'是否商务品': '非商务'})
    
    # 打印合并后的商务品数量
    print(f"合并后的商务品数量: {len(result[result['是否商务品'] == '商务']['证券代码'].unique())}")
    
    # 动态生成数值列
    numeric_cols = []
    for metric in ['关注人数', '持仓客户数', '保有金额']:
        current_col = f'{metric}（{current_str}）'
        previous_col = f'{metric}（{previous_str}）'
        
        # 验证列是否存在
        if current_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{current_col}'不存在，可用列：\n{available_cols}")
        if previous_col not in result.columns:
            available_cols = "\n".join(result.columns)
            raise KeyError(f"列名'{previous_col}'不存在，可用列：\n{available_cols}")
        
        numeric_cols.extend([current_col, previous_col])
        result[f'{metric}变动'] = result[current_col] - result[previous_col]

    # 数值类型转换
    result[numeric_cols] = result[numeric_cols].fillna(0).astype(int)
    
    # 计算商务品预期收入
    result['预期年收入(万元)'] = result.apply(
        lambda x: (x[f'保有金额（{current_str}）'] / 1e8) * (x['管理费率[单位]%'] / 100) * 0.4 * 10000 if x['是否商务品'] == '商务' else 0,
        axis=1
    )
    
    # 计算商务品占比
    result['商务品占比'] = result.apply(
        lambda x: 1 if x['是否商务品'] == '商务' else 0,
        axis=1
    )
    
    # 生成正确日期范围（MMDD-MMDD）
    date_range = f"{previous_date.strftime('%m%d')}-{current_date.strftime('%m%d')}"
    
    return result, date_range, current_str, previous_str

class BusinessETFAnalyzer:
    def __init__(self, data, date_range, current_date_str, previous_date_str):
        self.data = data
        self.date_range = date_range
        self.current_date = current_date_str
        self.previous_date = previous_date_str
        self.biz_data = data[data['是否商务品'] == '商务']
        self.output_dir = 'business_etf_analysis'
        
        # 创建输出目录
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        # 设置绘图风格
        sns.set(style="whitegrid")
        plt.rcParams['font.sans-serif'] = ['Arial Unicode MS']  # 设置中文字体
        plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题
    
    def analyze(self):
        """执行全面分析"""
        self.generate_summary_stats()
        self.analyze_by_company()
        self.analyze_by_index_type()
        self.analyze_by_index_level()
        self.analyze_performance()
        self.analyze_fee_distribution()
        self.analyze_customer_engagement()
        self.generate_report()
        
    def generate_summary_stats(self):
        """生成商务品总体统计"""
        print("生成商务品总体统计...")
        
        # 创建输出目录
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # 获取商务品数据
        self.biz_data = self.data[self.data['是否商务品'] == '商务'].copy()
        
        # 计算商务品总数
        biz_count = len(self.biz_data['证券代码'].unique())
        
        # 计算商务品总规模（亿元）
        biz_scale = self.biz_data[f'保有金额（{self.current_date}）'].sum() / 1e8
        
        # 计算所有ETF总规模（亿元）
        total_scale = self.data[f'保有金额（{self.current_date}）'].sum() / 1e8
        
        # 计算商务品占比
        biz_percentage = (biz_scale / total_scale) * 100 if total_scale > 0 else 0
        
        # 计算商务品总关注人数
        biz_attention = self.biz_data[f'关注人数（{self.current_date}）'].sum()
        
        # 计算所有ETF总关注人数
        total_attention = self.data[f'关注人数（{self.current_date}）'].sum()
        
        # 计算商务品关注人数占比
        attention_percentage = (biz_attention / total_attention) * 100 if total_attention > 0 else 0
        
        # 计算商务品总持仓客户数
        biz_holders = self.biz_data[f'持仓客户数（{self.current_date}）'].sum()
        
        # 计算所有ETF总持仓客户数
        total_holders = self.data[f'持仓客户数（{self.current_date}）'].sum()
        
        # 计算商务品持仓客户数占比
        holders_percentage = (biz_holders / total_holders) * 100 if total_holders > 0 else 0
        
        # 计算预期年收入
        # 检查是否有分成比例列
        if '分成比例' in self.biz_data.columns:
            # 使用实际分成比例计算
            expected_income = self.biz_data.apply(
                lambda x: (x[f'保有金额（{self.current_date}）'] / 1e8) * (x['管理费率[单位]%'] / 100) * x['分成比例'] * 10000,
                axis=1
            ).sum()
        else:
            # 使用默认分成比例0.4计算
            print("警告: 数据中缺少'分成比例'列，使用默认值0.4")
            expected_income = self.biz_data.apply(
                lambda x: (x[f'保有金额（{self.current_date}）'] / 1e8) * (x['管理费率[单位]%'] / 100) * 0.4 * 10000,
                axis=1
            ).sum()
        
        # 保存统计结果
        self.summary_stats = {
            '商务品总数': biz_count,
            '商务品总规模(亿元)': biz_scale,
            '商务品规模占比(%)': biz_percentage,
            '商务品总关注人数': biz_attention,
            '商务品关注人数占比(%)': attention_percentage,
            '商务品总持仓客户数': biz_holders,
            '商务品持仓客户数占比(%)': holders_percentage,
            '预期年收入(万元)': expected_income
        }
        
        # 创建商务品与非商务品对比数据
        biz_vs_nonbiz = pd.DataFrame([
            {
                '类型': '商务品',
                '数量': biz_count,
                '规模(亿元)': biz_scale,
                '关注人数': biz_attention,
                '持仓客户数': biz_holders
            },
            {
                '类型': '非商务品',
                '数量': len(self.data[self.data['是否商务品'] == '非商务']['证券代码'].unique()),
                '规模(亿元)': total_scale - biz_scale,
                '关注人数': total_attention - biz_attention,
                '持仓客户数': total_holders - biz_holders
            }
        ])
        
        # 绘制商务品与非商务品规模占比饼图
        plt.figure(figsize=(10, 6))
        plt.pie(
            [biz_scale, total_scale - biz_scale],
            labels=['商务品', '非商务品'],
            autopct='%1.1f%%',
            startangle=90,
            shadow=True
        )
        plt.axis('equal')
        plt.title('商务品与非商务品规模占比')
        plt.savefig(f'{self.output_dir}/商务品占比.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制商务品与非商务品客户参与度对比图
        plt.figure(figsize=(12, 8))
        x = np.arange(2)
        width = 0.35
        
        # 计算每个商务品和非商务品的平均关注人数和持仓客户数
        biz_avg_attention = biz_attention / biz_count if biz_count > 0 else 0
        nonbiz_count = len(self.data[self.data['是否商务品'] == '非商务']['证券代码'].unique())
        nonbiz_avg_attention = (total_attention - biz_attention) / nonbiz_count if nonbiz_count > 0 else 0
        
        biz_avg_holders = biz_holders / biz_count if biz_count > 0 else 0
        nonbiz_avg_holders = (total_holders - biz_holders) / nonbiz_count if nonbiz_count > 0 else 0
        
        plt.bar(x - width/2, [biz_avg_attention, biz_avg_holders], width, label='商务品平均')
        plt.bar(x + width/2, [nonbiz_avg_attention, nonbiz_avg_holders], width, label='非商务品平均')
        
        plt.xlabel('指标')
        plt.ylabel('人数')
        plt.title('商务品与非商务品客户参与度对比')
        plt.xticks(x, ['平均关注人数', '平均持仓客户数'])
        plt.legend()
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/商务品与非商务品客户参与度对比.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制商务品与非商务品平均变动对比图
        plt.figure(figsize=(12, 8))
        
        # 计算商务品和非商务品的平均变动
        biz_avg_attention_change = self.biz_data['关注人数变动'].mean()
        nonbiz_avg_attention_change = self.data[self.data['是否商务品'] == '非商务']['关注人数变动'].mean()
        
        biz_avg_holders_change = self.biz_data['持仓客户数变动'].mean()
        nonbiz_avg_holders_change = self.data[self.data['是否商务品'] == '非商务']['持仓客户数变动'].mean()
        
        biz_avg_amount_change = self.biz_data['保有金额变动'].mean() / 1e6  # 转换为百万元
        nonbiz_avg_amount_change = self.data[self.data['是否商务品'] == '非商务']['保有金额变动'].mean() / 1e6  # 转换为百万元
        
        x = np.arange(3)
        plt.bar(x - width/2, [biz_avg_attention_change, biz_avg_holders_change, biz_avg_amount_change], width, label='商务品平均')
        plt.bar(x + width/2, [nonbiz_avg_attention_change, nonbiz_avg_holders_change, nonbiz_avg_amount_change], width, label='非商务品平均')
        
        plt.xlabel('指标')
        plt.ylabel('变动值')
        plt.title('商务品与非商务品平均变动对比')
        plt.xticks(x, ['关注人数变动', '持仓客户数变动', '保有金额变动(百万元)'])
        plt.legend()
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/商务品与非商务品平均变动对比.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 基本统计
        total_count = len(self.biz_data['证券代码'].unique())
        total_scale = self.biz_data[f'保有金额（{self.current_date}）'].sum() / 1e8  # 转换为亿元
        all_etf_scale = self.data[f'保有金额（{self.current_date}）'].sum() / 1e8  # 转换为亿元
        percentage = (total_scale / all_etf_scale) * 100 if all_etf_scale > 0 else 0
        
        # 计算预期管理费收入
        total_expected_income = self.biz_data['预期年收入(万元)'].sum()
        
        # 计算平均管理费率
        avg_fee_rate = self.biz_data['管理费率[单位]%'].mean()
        
        # 计算加权平均管理费率（按保有金额加权）
        total_amount = self.biz_data[f'保有金额（{self.current_date}）'].sum()
        if total_amount > 0:
            weighted_fee_rate = (self.biz_data['管理费率[单位]%'] * self.biz_data[f'保有金额（{self.current_date}）']).sum() / total_amount
        else:
            weighted_fee_rate = 0
        
        # 计算平均分成比例 - 检查列是否存在
        avg_share_ratio = 0.4  # 默认值
        if '分成比例' in self.biz_data.columns:
            self.biz_data['分成比例'] = pd.to_numeric(self.biz_data['分成比例'], errors='coerce')
            avg_share_ratio = self.biz_data['分成比例'].mean()
            if pd.isna(avg_share_ratio):  # 如果结果是NaN，使用默认值
                avg_share_ratio = 0.4
        else:
            print("警告: 数据中缺少'分成比例'列，使用默认值0.4")
        
        # 保存统计结果
        self.summary_stats = {
            'total_count': total_count,
            'total_scale': total_scale,
            'all_etf_scale': all_etf_scale,
            'percentage': percentage,
            'total_expected_income': total_expected_income,
            'avg_fee_rate': avg_fee_rate,
            'weighted_fee_rate': weighted_fee_rate,  # 新增加权管理费率
            'avg_share_ratio': avg_share_ratio
        }
        
    def analyze_by_company(self):
        """按基金公司分析商务品"""
        print("按基金公司分析商务品...")
        
        # 按基金公司分组统计
        company_stats = self.biz_data.groupby('基金公司简称').agg({
            '证券代码': 'nunique',  # 商务品数量
            f'保有金额（{self.current_date}）': 'sum',  # 保有金额
            f'关注人数（{self.current_date}）': 'sum',  # 关注人数
            f'持仓客户数（{self.current_date}）': 'sum',  # 持仓客户数
            '管理费率[单位]%': 'mean',  # 平均管理费率
            '预期年收入(万元)': 'sum'  # 预期年收入
        }).reset_index()
        
        # 重命名列
        company_stats.columns = ['基金公司', '商务品数量', '保有金额', '关注人数', '持仓客户数', '平均管理费率', '预期年收入']
        
        # 转换保有金额为亿元
        company_stats['保有金额(亿元)'] = company_stats['保有金额'] / 1e8
        company_stats.drop('保有金额', axis=1, inplace=True)
        
        # 按保有金额排序
        company_stats = company_stats.sort_values('保有金额(亿元)', ascending=False)
        
        # 保存前20家公司的数据
        self.top_companies = company_stats.head(20)
        
        # 绘制前10家公司的商务品规模柱状图
        plt.figure(figsize=(12, 8))
        top10 = company_stats.head(10)
        sns.barplot(x='基金公司', y='保有金额(亿元)', data=top10)
        plt.title('前10家基金公司商务品规模')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/前10家公司商务品规模.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制前10家公司的商务品数量柱状图
        plt.figure(figsize=(12, 8))
        sns.barplot(x='基金公司', y='商务品数量', data=top10)
        plt.title('前10家基金公司商务品数量')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/前10家公司商务品数量.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制前10家公司的预期年收入柱状图
        plt.figure(figsize=(12, 8))
        sns.barplot(x='基金公司', y='预期年收入', data=top10)
        plt.title('前10家基金公司预期年收入(万元)')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        # 修改这里的文件名，确保与报告中引用的文件名一致
        plt.savefig(f'{self.output_dir}/前10家基金公司预期年收入.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 保存公司统计数据到CSV
        company_stats.to_csv(f'{self.output_dir}/基金公司商务品统计.csv', index=False, encoding='utf-8-sig')

    def analyze_by_index_type(self):
        """按指数类型分析商务品"""
        print("按指数类型分析商务品...")
        
        # 检查是否有分类列
        classification_columns = ['一级分类', '二级分类', '三级分类']
        available_columns = [col for col in classification_columns if col in self.biz_data.columns]
        
        if not available_columns:
            print("警告: 数据中缺少指数分类列，无法进行指数类型分析")
            return
        
        # 使用一级分类作为主要分类依据
        primary_classification = available_columns[0]
        print(f"使用 '{primary_classification}' 列进行指数类型分析")
        
        # 按指数类型分组统计
        index_type_stats = self.biz_data.groupby(primary_classification).agg({
            '证券代码': 'nunique',
            f'保有金额（{self.current_date}）': 'sum'
        }).reset_index()
        
        # 重命名列
        index_type_stats.columns = [primary_classification, '商务品数量', '保有金额']
        
        # 转换保有金额为亿元
        index_type_stats['保有金额(亿元)'] = index_type_stats['保有金额'] / 1e8
        
        # 计算占比
        total_amount = index_type_stats['保有金额(亿元)'].sum()
        index_type_stats['占比'] = index_type_stats['保有金额(亿元)'] / total_amount * 100 if total_amount > 0 else 0
        
        # 按保有金额排序
        index_type_stats = index_type_stats.sort_values('保有金额(亿元)', ascending=False)
        
        # 保存统计结果
        index_type_stats.to_csv(f'{self.output_dir}/商务品指数类型统计.csv', index=False, encoding='utf-8-sig')
        
        # 绘制指数类型规模分布饼图
        plt.figure(figsize=(12, 8))
        plt.pie(
            index_type_stats['保有金额(亿元)'],
            labels=index_type_stats[primary_classification],
            autopct='%1.1f%%',
            startangle=90,
            shadow=True
        )
        plt.axis('equal')
        plt.title('商务品指数类型规模分布')
        plt.savefig(f'{self.output_dir}/商务品指数类型规模分布.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制各指数类型商务品数量柱状图
        plt.figure(figsize=(12, 8))
        sns.barplot(x=primary_classification, y='商务品数量', data=index_type_stats)
        plt.title('各指数类型商务品数量')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数类型商务品数量.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制各指数类型商务品规模柱状图
        plt.figure(figsize=(12, 8))
        sns.barplot(x=primary_classification, y='保有金额(亿元)', data=index_type_stats)
        plt.title('各指数类型商务品规模(亿元)')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数类型商务品规模.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 保存指数类型统计数据
        self.index_type_stats = index_type_stats
        
        # 绘制指数类型商务品规模饼图（只取前8个类型，其余归为"其他"）
        plt.figure(figsize=(12, 10))
        top_types = index_type_stats.head(8)
        other_types = index_type_stats.iloc[8:] if len(index_type_stats) > 8 else pd.DataFrame()
        
        if len(other_types) > 0:
            other_amount = other_types['保有金额(亿元)'].sum()
            pie_data = pd.concat([
                top_types,
                pd.DataFrame([{
                    primary_classification: '其他',
                    '保有金额(亿元)': other_amount
                }])
            ])
        else:
            pie_data = top_types
        
        plt.pie(
            pie_data['保有金额(亿元)'],
            labels=pie_data[primary_classification],
            autopct='%1.1f%%',
            startangle=90,
            shadow=True
        )
        plt.axis('equal')
        plt.title('商务品指数类型规模分布(Top8)')
        plt.savefig(f'{self.output_dir}/商务品指数类型规模分布_Top8.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制指数类型商务品数量柱状图(Top10)
        plt.figure(figsize=(14, 8))
        sns.barplot(x=primary_classification, y='商务品数量', data=index_type_stats.head(10))
        plt.title('各指数类型商务品数量(Top10)')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数类型商务品数量_Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制指数类型商务品规模柱状图(Top10)
        plt.figure(figsize=(14, 8))
        sns.barplot(x=primary_classification, y='保有金额(亿元)', data=index_type_stats.head(10))
        plt.title('各指数类型商务品规模(亿元)(Top10)')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数类型商务品规模_Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 保存指数类型统计数据到CSV
        index_type_stats.to_csv(f'{self.output_dir}/指数类型商务品统计.csv', index=False, encoding='utf-8-sig')
    
    def analyze_by_index_level(self):
        """按指数层级分析商务品"""
        print("按指数层级分析商务品...")
        
        # 创建指数分级数据
        print("生成指数分级数据...")
        
        # 按跟踪指数代码分组
        index_groups = self.data.groupby('跟踪指数代码')
        
        index_levels = []
        for index_code, group in index_groups:
            try:
                # 获取最低费率基金
                lowest_fee_group = group[group['管理费率[单位]%'] == group['管理费率[单位]%'].min()]
                if len(lowest_fee_group) > 0:
                    # 如果有多个最低费率基金，选择交易量最大的
                    lowest_fee_fund = lowest_fee_group.loc[lowest_fee_group[f'关注人数（{self.current_date}）'].idxmax()]
                    is_lowest_fee_business = lowest_fee_fund['是否商务品'] == '商务'
                else:
                    is_lowest_fee_business = False
                
                # 获取最高交易量基金
                highest_volume_fund = group.loc[group[f'关注人数（{self.current_date}）'].idxmax()]
                is_highest_volume_business = highest_volume_fund['是否商务品'] == '商务'
                
                # 获取最大规模基金
                largest_scale_fund = group.loc[group[f'保有金额（{self.current_date}）'].idxmax()]
                is_largest_scale_business = largest_scale_fund['是否商务品'] == '商务'
                
                # 获取商务品数量
                business_count = len(group[group['是否商务品'] == '商务'])
                
                # 跟踪指数分级逻辑
                if is_highest_volume_business and is_lowest_fee_business and is_largest_scale_business:
                    index_level = '一级'
                elif is_highest_volume_business and is_lowest_fee_business and not is_largest_scale_business:
                    index_level = '二级'
                elif is_highest_volume_business and is_largest_scale_business and not is_lowest_fee_business:
                    index_level = '三级'
                elif is_lowest_fee_business and is_largest_scale_business and not is_highest_volume_business:
                    index_level = '四级'
                elif is_highest_volume_business:
                    index_level = '五级'
                elif is_largest_scale_business:
                    index_level = '六级'
                elif is_lowest_fee_business:
                    index_level = '七级'
                elif business_count > 0:
                    index_level = '八级'
                else:
                    index_level = '九级'
                
                # 计算指数总规模
                total_scale = group[f'保有金额（{self.current_date}）'].sum() / 1e8
                
                # 计算商务品规模
                business_scale = group[group['是否商务品'] == '商务'][f'保有金额（{self.current_date}）'].sum() / 1e8
                
                # 添加到结果列表
                index_levels.append({
                    '跟踪指数代码': index_code,
                    '跟踪指数名称': group['跟踪指数名称'].iloc[0] if '跟踪指数名称' in group.columns else '未知',
                    '指数层级': index_level,
                    'ETF数量': len(group),
                    '商务品数量': business_count,
                    '商务品规模(亿元)': business_scale,
                    '总规模(亿元)': total_scale,
                    '商务品占比(%)': (business_scale / total_scale * 100) if total_scale > 0 else 0
                })
            except Exception as e:
                print(f"处理指数 {index_code} 时出错: {str(e)}")
                continue
        
        # 创建DataFrame
        index_level_stats = pd.DataFrame(index_levels)
        
        # 按指数层级分组统计
        level_summary = index_level_stats.groupby('指数层级').agg({
            '跟踪指数代码': 'count',
            'ETF数量': 'sum',
            '商务品数量': 'sum',
            '商务品规模(亿元)': 'sum',
            '总规模(亿元)': 'sum'
        }).reset_index()
        
        # 重命名列
        level_summary.columns = ['指数层级', '指数数量', 'ETF数量', '商务品数量', '商务品规模(亿元)', '总规模(亿元)']
        
        # 计算商务品占比
        level_summary['商务品占比(%)'] = level_summary['商务品规模(亿元)'] / level_summary['总规模(亿元)'] * 100
        
        # 按指数层级排序（一级到九级）
        level_order = {'一级': 1, '二级': 2, '三级': 3, '四级': 4, '五级': 5, '六级': 6, '七级': 7, '八级': 8, '九级': 9}
        level_summary['排序'] = level_summary['指数层级'].map(level_order)
        level_summary = level_summary.sort_values('排序').drop('排序', axis=1)
        
        # 保存指数层级统计数据
        # 保存指数层级统计数据
        self.index_level_stats = level_summary
        index_level_stats.to_csv(f'{self.output_dir}/指数层级详细数据.csv', index=False, encoding='utf-8-sig')
        level_summary.to_csv(f'{self.output_dir}/指数层级统计.csv', index=False, encoding='utf-8-sig')
        
        # 绘制指数层级分布饼图
        plt.figure(figsize=(12, 8))
        plt.pie(
            level_summary['指数数量'],
            labels=level_summary['指数层级'],
            autopct='%1.1f%%',
            startangle=90,
            shadow=True
        )
        plt.axis('equal')
        plt.title('跟踪指数层级分布')
        plt.savefig(f'{self.output_dir}/跟踪指数层级分布.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 添加商务品指数层级规模分布图 (与报告中引用的文件名一致)
        plt.figure(figsize=(12, 8))
        plt.pie(
            level_summary['商务品规模(亿元)'],
            labels=level_summary['指数层级'],
            autopct='%1.1f%%',
            startangle=90,
            shadow=True
        )
        plt.axis('equal')
        plt.title('商务品指数层级规模分布')
        plt.savefig(f'{self.output_dir}/商务品指数层级规模分布.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制各层级商务品规模柱状图 - 修改文件名
        plt.figure(figsize=(12, 8))
        sns.barplot(x='指数层级', y='商务品规模(亿元)', data=level_summary)
        plt.title('各层级商务品规模(亿元)')
        plt.xticks(rotation=0)
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数层级商务品规模.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制各层级商务品数量柱状图 - 修改文件名
        plt.figure(figsize=(12, 8))
        sns.barplot(x='指数层级', y='商务品数量', data=level_summary)
        plt.title('各层级商务品数量')
        plt.xticks(rotation=0)
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各指数层级商务品数量.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制各层级指数数量柱状图
        plt.figure(figsize=(12, 8))
        sns.barplot(x='指数层级', y='指数数量', data=level_summary)
        plt.title('各层级指数数量')
        plt.xticks(rotation=0)
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/各层级指数数量.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"指数层级分析完成，共分析 {len(index_levels)} 个跟踪指数")
        
    def analyze_performance(self):
        """分析商务品业绩表现"""
        print("分析商务品业绩表现...")
        
        # 检查是否有必要的列
        required_cols = ['关注人数变动', '持仓客户数变动', '保有金额变动']
        for col in required_cols:
            if col not in self.biz_data.columns:
                print(f"警告: 数据中缺少'{col}'列，无法进行业绩表现分析")
                return
        
        # 按变动幅度排序
        attention_top = self.biz_data.sort_values('关注人数变动', ascending=False).head(10)
        holders_top = self.biz_data.sort_values('持仓客户数变动', ascending=False).head(10)
        amount_top = self.biz_data.sort_values('保有金额变动', ascending=False).head(10)
        
        # 绘制关注人数变动Top10
        plt.figure(figsize=(12, 8))
        sns.barplot(x='关注人数变动', y='产品名称', data=attention_top)
        plt.title('关注人数变动Top10商务品')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/关注人数变动Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制持仓客户数变动Top10
        plt.figure(figsize=(12, 8))
        sns.barplot(x='持仓客户数变动', y='产品名称', data=holders_top)
        plt.title('持仓客户数变动Top10商务品')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/持仓客户数变动Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制保有金额变动Top10（百万元）
        plt.figure(figsize=(12, 8))
        amount_top['保有金额变动(百万)'] = amount_top['保有金额变动'] / 1e6
        sns.barplot(x='保有金额变动(百万)', y='产品名称', data=amount_top)
        plt.title('保有金额变动Top10商务品(百万元)')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/保有金额变动Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 保存业绩表现数据
        performance_data = pd.DataFrame({
            '关注人数变动Top10': attention_top['产品名称'].values,
            '持仓客户数变动Top10': holders_top['产品名称'].values,
            '保有金额变动Top10': amount_top['产品名称'].values
        })
        performance_data.to_csv(f'{self.output_dir}/商务品业绩表现Top10.csv', index=False, encoding='utf-8-sig')
    
    def analyze_fee_distribution(self):
        """分析商务品管理费率分布"""
        print("分析商务品管理费率分布...")
        
        # 检查是否有管理费率列
        if '管理费率[单位]%' not in self.biz_data.columns:
            print("警告: 数据中缺少'管理费率[单位]%'列，无法进行管理费率分析")
            return
        
        # 管理费率分布
        plt.figure(figsize=(12, 8))
        sns.histplot(self.biz_data['管理费率[单位]%'], bins=20, kde=True)
        plt.title('商务品管理费率分布')
        plt.xlabel('管理费率(%)')
        plt.ylabel('数量')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/商务品管理费率分布.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 管理费率与规模的关系
        plt.figure(figsize=(12, 8))
        self.biz_data['保有金额(亿元)'] = self.biz_data[f'保有金额（{self.current_date}）'] / 1e8
        sns.scatterplot(x='管理费率[单位]%', y='保有金额(亿元)', data=self.biz_data)
        plt.title('商务品管理费率与规模关系')
        plt.xlabel('管理费率(%)')
        plt.ylabel('保有金额(亿元)')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/商务品管理费率与规模关系.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 计算加权平均管理费率
        total_amount = self.biz_data[f'保有金额（{self.current_date}）'].sum()
        if total_amount > 0:
            weighted_fee_rate = (self.biz_data['管理费率[单位]%'] * self.biz_data[f'保有金额（{self.current_date}）']).sum() / total_amount
        else:
            weighted_fee_rate = 0
        
        # 管理费率统计
        fee_stats = {
            '平均管理费率(%)': self.biz_data['管理费率[单位]%'].mean(),
            '加权平均管理费率(%)': weighted_fee_rate,  # 新增加权管理费率
            '中位数管理费率(%)': self.biz_data['管理费率[单位]%'].median(),
            '最高管理费率(%)': self.biz_data['管理费率[单位]%'].max(),
            '最低管理费率(%)': self.biz_data['管理费率[单位]%'].min(),
            '标准差': self.biz_data['管理费率[单位]%'].std()
        }
        
        # 保存管理费率统计数据
        pd.DataFrame([fee_stats]).to_csv(f'{self.output_dir}/商务品管理费率统计.csv', index=False, encoding='utf-8-sig')
    
    def analyze_customer_engagement(self):
        """分析商务品客户参与度"""
        print("分析商务品客户参与度...")
        
        # 计算每个商务品的平均关注人数和持仓客户数
        self.biz_data['平均关注人数'] = self.biz_data[f'关注人数（{self.current_date}）']
        self.biz_data['平均持仓客户数'] = self.biz_data[f'持仓客户数（{self.current_date}）']
        
        # 关注人数Top10
        attention_top = self.biz_data.sort_values('平均关注人数', ascending=False).head(10)
        
        # 持仓客户数Top10
        holders_top = self.biz_data.sort_values('平均持仓客户数', ascending=False).head(10)
        
        # 绘制关注人数Top10
        plt.figure(figsize=(12, 8))
        sns.barplot(x='平均关注人数', y='产品名称', data=attention_top)
        plt.title('关注人数Top10商务品')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/关注人数Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 绘制持仓客户数Top10
        plt.figure(figsize=(12, 8))
        sns.barplot(x='平均持仓客户数', y='产品名称', data=holders_top)
        plt.title('持仓客户数Top10商务品')
        plt.tight_layout()
        plt.savefig(f'{self.output_dir}/持仓客户数Top10.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # 保存客户参与度数据
        engagement_data = pd.DataFrame({
            '关注人数Top10': attention_top['产品名称'].values,
            '持仓客户数Top10': holders_top['产品名称'].values
        })
        engagement_data.to_csv(f'{self.output_dir}/商务品客户参与度Top10.csv', index=False, encoding='utf-8-sig')
    
    def generate_report(self):
        """生成商务品分析报告"""
        print("生成商务品分析报告...")
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档标题
        doc.add_heading('ETF商务品分析报告', 0)
        
        # 添加日期范围
        doc.add_paragraph(f'数据日期范围: {self.date_range}')
        
        # 添加基本统计信息
        doc.add_heading('1. 商务品基本情况', 1)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        # 创建中文指标名称映射
        stats_mapping = {
            'total_count': '商务品总数',
            'total_scale': '商务品总规模(亿元)',
            'all_etf_scale': 'ETF总规模(亿元)',
            'percentage': '商务品规模占比(%)',
            'total_expected_income': '预期年收入(万元)',
            'avg_fee_rate': '平均管理费率(%)',
            'weighted_fee_rate': '加权平均管理费率(%)',  # 新增加权管理费率
            'avg_share_ratio': '平均分成比例'
        }
        
        # 添加统计数据行
        for key, value in self.summary_stats.items():
            row_cells = stats_table.add_row().cells
            # 使用中文指标名称
            row_cells[0].text = stats_mapping.get(key, key)
            
            if isinstance(value, (int, float)):
                # 根据不同类型的指标格式化数值
                if key in ['total_scale', 'all_etf_scale', 'total_expected_income', 'avg_fee_rate', 'weighted_fee_rate', 'avg_share_ratio']:
                    # 显示小数点后两位
                    row_cells[1].text = f"{value:.2f}"
                    # 添加单位
                    if key == 'total_scale' or key == 'all_etf_scale':
                        row_cells[1].text += " 亿元"
                    elif key == 'total_expected_income':
                        row_cells[1].text += " 万元"
                    elif key == 'avg_fee_rate' or key == 'weighted_fee_rate':
                        row_cells[1].text += "%"
                elif key == 'percentage':
                    row_cells[1].text = f"{value:.2f}%"
                else:
                    row_cells[1].text = f"{value:,}"
            else:
                row_cells[1].text = str(value)
        
        # 添加商务品占比图
        doc.add_paragraph('商务品与非商务品规模占比:')
        doc.add_picture(f'{self.output_dir}/商务品占比.png', width=Inches(6))
        
        # 添加客户参与度对比图
        doc.add_paragraph('商务品与非商务品客户参与度对比:')
        doc.add_picture(f'{self.output_dir}/商务品与非商务品客户参与度对比.png', width=Inches(6))
        
        # 添加平均变动对比图
        doc.add_paragraph('商务品与非商务品平均变动对比:')
        doc.add_picture(f'{self.output_dir}/商务品与非商务品平均变动对比.png', width=Inches(6))
        
        # 添加基金公司分析
        doc.add_heading('2. 基金公司分析', 1)
        doc.add_paragraph('前10家基金公司商务品规模:')
        doc.add_picture(f'{self.output_dir}/前10家公司商务品规模.png', width=Inches(6))
        
        doc.add_paragraph('前10家基金公司商务品数量:')
        doc.add_picture(f'{self.output_dir}/前10家公司商务品数量.png', width=Inches(6))
        
        doc.add_paragraph('前10家基金公司预期年收入:')
        doc.add_picture(f'{self.output_dir}/前10家基金公司预期年收入.png', width=Inches(6))
        
        # 添加指数类型分析（如果有）
        if hasattr(self, 'index_type_stats'):
            doc.add_heading('3. 指数类型分析', 1)
            doc.add_paragraph('商务品指数类型规模分布:')
            doc.add_picture(f'{self.output_dir}/商务品指数类型规模分布.png', width=Inches(6))
            
            doc.add_paragraph('各指数类型商务品数量:')
            doc.add_picture(f'{self.output_dir}/各指数类型商务品数量.png', width=Inches(6))
            
            doc.add_paragraph('各指数类型商务品规模:')
            doc.add_picture(f'{self.output_dir}/各指数类型商务品规模.png', width=Inches(6))
        
        # 添加指数层级分析（如果有）
        if hasattr(self, 'index_level_stats'):
            doc.add_heading('4. 指数层级分析', 1)
            doc.add_paragraph('商务品指数层级规模分布:')
            doc.add_picture(f'{self.output_dir}/商务品指数层级规模分布.png', width=Inches(6))
            
            doc.add_paragraph('各层级商务品数量:')
            doc.add_picture(f'{self.output_dir}/各指数层级商务品数量.png', width=Inches(6))
            
            doc.add_paragraph('各层级商务品规模:')
            doc.add_picture(f'{self.output_dir}/各指数层级商务品规模.png', width=Inches(6))
        
        # 保存文档
        report_path = f'{self.output_dir}/ETF商务品分析报告_{self.date_range}.docx'
        doc.save(report_path)
        print(f"报告已保存至: {report_path}")


# 在文件末尾添加以下代码
if __name__ == "__main__":
    show_version()
    print("开始加载数据...")
    try:
        data, date_range, current_date, previous_date = load_data()
        print(f"数据加载完成，日期范围: {date_range}")
        
        analyzer = BusinessETFAnalyzer(data, date_range, current_date, previous_date)
        analyzer.analyze()
        print("分析完成！")
    except Exception as e:
        print(f"错误: {str(e)}")